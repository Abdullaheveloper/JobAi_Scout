"""Generate a reader-friendly explanation of the JobAI Scout Voice Assistant."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name("JobAI_Scout_Voice_Assistant_Explained.docx")
NAVY, INDIGO, VIOLET, SLATE = "0B1028", "4F46E5", "7C3AED", "334155"


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), color)
    props.append(node)


def line_below(paragraph):
    props = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), INDIGO)
    border.append(bottom)
    props.append(border)


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Cm(1.7)
section.left_margin = section.right_margin = Cm(1.8)
normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Aptos", Pt(10)
normal.font.color.rgb = RGBColor.from_string("172033")
normal.paragraph_format.space_after = Pt(6)
for name, size, color in [("Heading 1", 17, NAVY), ("Heading 2", 13, INDIGO), ("Heading 3", 11, VIOLET)]:
    style = doc.styles[name]
    style.font.name, style.font.size, style.font.bold = "Aptos Display", Pt(size), True
    style.font.color.rgb = RGBColor.from_string(color)


def h(text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        line_below(p)
    return p


def p(text):
    return doc.add_paragraph(text)


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def info(label, text, fill="EEF2FF"):
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    shade(c, fill)
    para = c.paragraphs[0]
    para.add_run(label + " ").bold = True
    para.add_run(text)
    doc.add_paragraph()


def grid(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, text in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = text
        shade(c, NAVY)
        for run in c.paragraphs[0].runs:
            run.font.bold, run.font.color.rgb = True, RGBColor(255, 255, 255)
    for index, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if index % 2 == 0:
                shade(cells[i], "F8FAFC")
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()


def flow(steps):
    t = doc.add_table(rows=1, cols=len(steps))
    for i, text in enumerate(steps):
        cell = t.cell(0, i)
        shade(cell, "EEF2FF" if i % 2 == 0 else "EDE9FE")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold, run.font.size, run.font.color.rgb = True, Pt(8), RGBColor.from_string(NAVY)
    doc.add_paragraph()


cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("How the Voice Assistant Works in JobAI Scout")
r.font.name, r.font.size, r.font.bold, r.font.color.rgb = "Aptos Display", Pt(25), True, RGBColor.from_string(NAVY)
subtitle = doc.add_paragraph("A code-based explanation of microphone capture, recognition, AI responses, speech playback, history and safety")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.color.rgb, subtitle.runs[0].font.size = RGBColor.from_string(SLATE), Pt(11)
doc.add_paragraph()
info("Purpose:", "This document explains the current Voice Assistant implementation. It describes what the user experiences, what runs inside the browser, which protected functions process the question, what data may be stored, and which security improvements are still required before production launch.")

h("1. The feature in one sentence")
p("The Voice Assistant lets an authenticated job seeker ask a career question by speaking. The browser records the microphone, uses browser speech recognition when available, falls back to server transcription when necessary, sends the recognized text to the AI career assistant, displays the answer in the current conversation, and plays a spoken response when text-to-speech succeeds.")
flow(["Start assistant", "Capture microphone", "Recognize / transcribe", "AI career answer", "Display response", "Play voice", "Keep private history"])
grid(["Layer", "Main responsibility", "Key code"], [
    ["Voice UI", "State-aware interface, controls, waveform, transcript, current-session messages and error guidance.", "src/components/voice/VoiceMode.tsx"],
    ["Browser capture", "MediaRecorder, getUserMedia, AudioContext analyser, silence timer and Web Speech API wrapper.", "VoiceMode.tsx and src/lib/voice/recognition.ts"],
    ["Transcription", "Converts recorded audio to text only when browser recognition did not produce usable text.", "supabase/functions/voice-transcribe"],
    ["AI conversation", "Authenticates user, optionally retrieves private document context, creates an answer and persists conversation records.", "supabase/functions/voice-chat"],
    ["Speech output", "Converts final answer text into generated audio; browser speech synthesis is the fallback.", "supabase/functions/voice-tts"],
    ["Private data", "Stores conversations/messages and optional user/assistant audio under user-scoped records or paths.", "voice_conversations, voice_messages, voice-history storage"],
], [3.2, 7.2, 6.2])

h("2. The user interface and its states")
p("VoiceMode does not treat voice capture as one single button action. It maintains an explicit status value so the interface can explain what the system is doing. The screen uses Framer Motion for gentle state transitions, a waveform while audio is active, an AI-thinking indicator while requests are running, and a live transcript card while the user is speaking.")
grid(["State", "Meaning for the user", "What the system is doing"], [
    ["idle", "Ready to start a new voice turn.", "No active recorder, request or playback."],
    ["permission", "Preparing microphone access.", "Checks secure context, browser support and permission state before capture."],
    ["listening", "The assistant is waiting for the first speech.", "MediaRecorder, audio analyser and optional browser recognition are active."],
    ["speaking-user", "Speech has been detected.", "Resets no-speech and silence timers; updates waveform and live transcript."],
    ["silence", "The user has paused after speaking.", "Runs the three-second end-of-turn timer."],
    ["uploading", "The question is being understood.", "Uses saved browser text or sends recorded audio to transcription."],
    ["thinking", "JobAI Scout is preparing an answer.", "Calls voice-chat and waits for the AI result."],
    ["speaking", "The assistant is reading the answer.", "Plays generated audio or browser speech-synthesis fallback."],
    ["paused / ended", "The previous response or entire session has stopped.", "Resources are stopped; Start Assistant is available again."],
    ["error", "A friendly recovery message is shown.", "A microphone, recognition, network or no-speech condition requires user action."],
], [3.0, 5.5, 8.1])

h("3. The buttons are intentionally separate")
grid(["Control", "Exact behavior"], [
    ["Start Assistant", "Requests microphone permission, then begins a new listening session. It is disabled while recording, thinking or speaking to prevent duplicate sessions."],
    ["Stop Listening", "Stops only the active recording. It requests the recorder’s final buffered data, preserves final and interim browser transcript text, then processes the question."],
    ["Stop Speaking", "Cancels browser speech synthesis and pauses any generated Audio object. It does not end the entire assistant session."],
    ["End Assistant", "Aborts in-flight network requests, discards an active recording, stops microphone tracks, playback, timers and animation analysis, then clears active voice state."],
    ["Try Again", "Appears after a friendly error and starts the microphone flow again after the user has corrected the issue."],
], [4.2, 12.4])
info("Why this matters:", "Using separate controls avoids a confusing multi-purpose microphone button. A user can stop capture, interrupt spoken audio, or end the entire session without guessing which side effect will occur.", "DCFCE7")

h("4. Starting the microphone safely")
p("When Start Assistant is pressed, the component first prevents concurrent startup with isStarting and recorder/request checks. It stops current playback, sets the UI to permission, and verifies that the page is a secure context. Browsers allow microphone access on HTTPS sites and localhost, but normally block it on ordinary HTTP network addresses.")
grid(["Check", "Reason", "User-facing response"], [
    ["window.isSecureContext", "Microphone APIs require HTTPS or localhost.", "Secure connection required guidance."],
    ["navigator.mediaDevices and MediaRecorder", "The browser must support capture APIs.", "Voice capture unavailable guidance."],
    ["AudioContext", "Required for the audio analyser used by the waveform and silence detection.", "Unsupported browser guidance."],
    ["Permissions API when available", "Detects a previously denied microphone grant before requesting it again.", "Allow microphone access guidance."],
    ["getUserMedia", "Asks the browser/operating system for the actual microphone stream.", "Maps NotAllowed, NotFound, NotReadable and related errors to friendly messages."],
], [4.2, 7.2, 5.2])
p("The microphone is requested with echoCancellation, noiseSuppression and autoGainControl. These browser-level constraints are intended to make human speech clearer and reduce accidental feedback/noise. Their exact behavior depends on the device and browser.")

h("5. Recording, waveform and silence detection")
p("After getUserMedia succeeds, the code uses the same microphone stream in three ways: MediaRecorder stores chunks of audio; VoiceRecognition can produce browser text; and an AudioContext analyser measures the current signal level for the waveform and end-of-speech behavior.")
flow(["Microphone stream", "MediaRecorder chunks", "Web Speech interim/final text", "Audio analyser level", "3-second silence", "onstop creates Blob", "Process one question"])
grid(["Mechanism", "Current implementation"], [
    ["MediaRecorder", "Begins with a 250 ms timeslice. Each non-empty dataavailable event is appended to chunks."],
    ["Waveform", "Analyser data is sampled through requestAnimationFrame. The computed level updates WaveformVisualizer; reduced-motion users receive a calmer view."],
    ["Initial wait", "If no speech signal is found, a 10-second no-speech timer stops the recorder rather than ending immediately."],
    ["Speech threshold", "A quiet-microphone threshold of 0.012 marks that the user has begun speaking."],
    ["End-of-turn rule", "After speech is detected, absence of signal starts a three-second silence timer. Speech again clears the timer."],
    ["Manual stop", "Stop Listening calls requestData before recorder.stop so final buffered audio is more likely to be included."],
], [4.2, 12.4])
info("Important behavior:", "The three-second rule begins only after the system has detected speech. It is not intended to stop the user merely because they paused before starting their question.")

h("6. Browser speech recognition and transcription fallback")
p("The VoiceRecognition wrapper uses the browser’s SpeechRecognition or webkitSpeechRecognition API when it exists. It runs continuously with interim results. Each result is marked either final or interim. VoiceMode stores final transcript text separately from the latest interim text. This is important because a user may press Stop Listening before the browser has promoted its last words to final text.")
grid(["Path", "When it is used", "Result"], [
    ["Browser recognition", "Supported browser returns a usable transcript.", "The browser text is used immediately; this avoids an extra server transcription request."],
    ["Interim preservation", "The user stops listening while words are still interim.", "Final plus interim text are combined and submitted instead of being discarded."],
    ["Server fallback", "No usable browser transcript exists but there is a usable recording.", "voice-transcribe receives multipart audio and asks Gemini to return only spoken text."],
    ["No speech", "No transcript and recording is smaller than the minimum threshold, or transcription returns empty text.", "The UI shows a friendly no-speech message and Try Again action."],
], [4.0, 5.2, 7.4])
p("The transcription function sends an audio file as Gemini inline data with an optional expected language. It expects plain text and returns a 422 error when the model finds no speech. The main client then maps errors into a user-friendly recognition, network or no-speech category.")

h("7. The AI question and answer pipeline")
p("After text exists, processRecording appends the user question to the visible current-session conversation and begins a non-blocking upload of the original recording to private voice-history storage. It then calls voice-chat. This Edge Function is the main protected intelligence boundary: it checks the bearer token, obtains the authenticated user, creates or continues a conversation, retrieves relevant private context when configured, and generates a career-focused answer.")
flow(["Recognized text", "voice-chat bearer token", "Validate current user", "Load/create conversation", "Retrieve private knowledge context", "Generate answer", "Persist messages", "Return answer + conversation ID"])
grid(["voice-chat responsibility", "Explanation"], [
    ["Authentication", "Creates a Supabase client with the request Authorization header and calls auth.getUser. It returns 401 when no authenticated user is present."],
    ["Input safety", "Rejects empty questions and applies basic prompt-injection pattern removal before AI processing."],
    ["Language / personality", "Detects language patterns and supports a selected personality mode for the system prompt."],
    ["Private context", "Uses user-owned knowledge-base chunks and semantic embeddings to retrieve relevant document content for a question."],
    ["Conversation continuity", "Accepts an existing conversation ID or creates one, then persists user and assistant message records."],
    ["Answer", "Uses configured Gemini generation helpers to provide career-oriented, voice-friendly guidance."],
], [4.4, 12.2])

h("8. Text-to-speech and playback")
p("Once voice-chat returns an answer, the client appends it to the conversation panel and calls voice-tts. The generated audio is played with an HTML Audio object. After playback begins, the audio is uploaded to private voice-history storage and the most recent assistant message is updated with its audio path. This background storage design means history work should not delay the first audible response.")
grid(["Playback stage", "Behavior"], [
    ["Primary TTS", "voice-tts asks the configured Gemini TTS model to read the final answer naturally and returns WAV audio."],
    ["Playback", "The browser constructs an Audio object from a temporary object URL and changes the UI to speaking."],
    ["Stop Speaking", "Pauses the Audio object, rewinds it, cancels browser synthesis and changes speaking state to paused."],
    ["Fallback", "If the server TTS call fails but the answer exists, SpeechSynthesisUtterance reads the text using a browser-provided voice."],
    ["History replay", "When a history item has an audio_path, the UI asks Storage for a signed URL valid for one hour before playback."],
], [4.2, 12.4])
info("Current implementation detail:", "VoiceMode sends voice 'Eve' and a speed value, while the current voice-tts function accepts a limited Gemini voice list and does not use the supplied speed. The function falls back to Kore for an unsupported voice. Align the UI options and API contract before presenting voice selection/speed as guaranteed server-TTS behavior.", "FEF3C7")

h("9. Conversation history and privacy")
grid(["Data", "Where it is used", "Privacy expectation"], [
    ["Current session messages", "React state in VoiceMode, shown in the Conversation panel.", "Exists in the browser for the active page/session."],
    ["Conversation and message text", "voice_conversations and voice_messages records created by voice-chat.", "Must be scoped by authenticated user and protected by database policies."],
    ["User recording", "Optional upload to voice-history under user ID and a unique path.", "Original audio is private account data, not a public asset."],
    ["Assistant audio", "Uploaded after playback begins; its path is attached to the assistant history message.", "Replay uses a short-lived signed URL, not a permanently public URL."],
    ["Knowledge-base document", "Uploaded through the assistant document button and indexed for private retrieval.", "Should only be retrieved for the owning authenticated user."],
], [4.0, 5.8, 6.8])
p("The permission dialog correctly explains that audio may be retained privately in the account so a user can replay it later. This is more accurate than claiming that voice data is never stored. Product privacy policy and retention controls should stay consistent with this behavior.")

h("10. Cleanup and duplicate-request protection")
p("Voice features can leak microphone access, timers and audio resources if they are not closed carefully. VoiceMode uses refs rather than repeated state objects for the recorder, stream, analyser context, timers, recognition instance, abort controller and player. This allows explicit cleanup even while the UI is re-rendering.")
grid(["Protection", "How it works"], [
    ["One active start", "isStarting, recorder state and requestAbort checks block rapid repeated Start clicks."],
    ["One request", "processRecording returns if an AbortController already represents an in-flight request."],
    ["Timer cleanup", "clearTimers clears silence and no-speech timers when recording ends or session ends."],
    ["Media cleanup", "stopRecording stops recognition, flushes/stops recorder, stops every MediaStream track and closes the AudioContext analyser."],
    ["Animation cleanup", "stopAnalysis cancels requestAnimationFrame before closing the analyser context."],
    ["Network cleanup", "End Assistant aborts the current fetch controller so a late response cannot continue the active session."],
    ["Unmount cleanup", "A React effect calls endAssistant when VoiceMode is removed from the page."],
], [4.5, 12.1])

h("11. Error handling")
grid(["Problem", "Friendly message path"], [
    ["Permission denied", "Permission API or NotAllowedError maps to guidance telling the user to allow microphone access in browser site settings."],
    ["HTTP / insecure origin", "A failed secure-context check maps to a message requiring HTTPS or localhost."],
    ["No microphone / busy device", "NotFound, Overconstrained, NotReadable and Abort errors map to unavailable or busy microphone guidance."],
    ["Unsupported browser", "Missing getUserMedia, MediaRecorder or AudioContext maps to browser support guidance."],
    ["No speech", "Very small recording, empty recognition result or empty transcription maps to a retry message."],
    ["Network / AI failure", "Fetch-related errors map to connection guidance; other failures map to recognition/assistant retry guidance."],
], [4.6, 12.0])

h("12. Production security improvements required")
p("The user experience and voice-chat authorization are strong foundations, but the following backend items should be addressed before public production use. These are security requirements, not only UI improvements.")
grid(["Finding", "Why it matters", "Required improvement"], [
    ["voice-transcribe has no token validation", "Any internet client may be able to call the endpoint and consume paid Gemini transcription capacity.", "Require Authorization, validate auth.getUser, reject unauthenticated requests and add rate limits."],
    ["voice-tts has no token validation", "Any internet client may be able to consume TTS capacity through the server’s Gemini key.", "Require verified authenticated user, enforce request limits and consider per-user quotas."],
    ["CORS allows every origin", "Broad cross-origin access makes unwanted browser use easier.", "Restrict allowed origins to the production domain and approved development origins."],
    ["TTS contract mismatch", "Client voice/speed controls do not exactly match server support; history labels use MP3 while the function returns WAV.", "Use one shared request schema, validate supported voice/speed values and store correct audio content type/extension."],
    ["Voice storage / RLS review", "Audio and transcripts are sensitive personal data.", "Audit storage and table RLS policies with two test accounts; document retention and deletion behavior."],
], [4.3, 6.0, 6.3])
info("Secret rule:", "GEMINI_API_KEY and SUPABASE_SERVICE_ROLE_KEY belong only in Supabase Edge Function secrets. They must never be added to Vite VITE variables, browser code, GitHub, screenshots or the browser extension.", "FEE2E2")

h("13. Reader summary")
p("JobAI Scout Voice Assistant is a multi-stage system. The browser owns microphone permission, recording, live visual feedback and direct user controls. Browser recognition gives the fastest text when available; server transcription is a fallback. A protected chat function turns the question into a career answer with user-scoped context, then TTS plays the response. The implementation intentionally cleans up media, timers, animation and requests to prevent duplicate sessions. Before public release, secure every voice endpoint with token validation, align the TTS contract, and complete a privacy/RLS audit for stored voice data.")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("JobAI Scout Voice Assistant - Technical Explanation")
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor.from_string(SLATE)
doc.core_properties.title = "How the Voice Assistant Works in JobAI Scout"
doc.core_properties.author = "JobAI Scout"
doc.save(OUT)
print(OUT)
