"""Generate a reader-friendly explanation of the JobAI Scout CV upload feature."""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name("JobAI_Scout_CV_Upload_Explained.docx")
NAVY, INDIGO, VIOLET, SLATE = "0B1028", "4F46E5", "7C3AED", "334155"


def shade(cell, value):
    props = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), value)
    props.append(fill)


def bottom_border(paragraph):
    props = paragraph._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    line = OxmlElement("w:bottom")
    line.set(qn("w:val"), "single")
    line.set(qn("w:sz"), "8")
    line.set(qn("w:color"), INDIGO)
    bdr.append(line)
    props.append(bdr)


doc = Document()
section = doc.sections[0]
for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(section, attr, Cm(1.7 if "margin" in attr else 1.7))

normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Aptos", Pt(10)
normal.font.color.rgb = RGBColor.from_string("172033")
normal.paragraph_format.space_after = Pt(6)
for name, size, color in [("Heading 1", 17, NAVY), ("Heading 2", 13, INDIGO), ("Heading 3", 11, VIOLET)]:
    style = doc.styles[name]
    style.font.name, style.font.size, style.font.bold = "Aptos Display", Pt(size), True
    style.font.color.rgb = RGBColor.from_string(color)


def heading(text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        bottom_border(p)
    return p


def para(text):
    return doc.add_paragraph(text)


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def info(label, text, color="EEF2FF"):
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    shade(c, color)
    p = c.paragraphs[0]
    p.add_run(label + " ").bold = True
    p.add_run(text)
    doc.add_paragraph()


def grid(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, text in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = text
        shade(c, NAVY)
        for r in c.paragraphs[0].runs:
            r.font.bold, r.font.color.rgb = True, RGBColor(255, 255, 255)
    for n, row in enumerate(rows, start=1):
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if n % 2 == 0:
                shade(cells[i], "F8FAFC")
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()


def flow(steps):
    t = doc.add_table(rows=1, cols=len(steps))
    for i, text in enumerate(steps):
        c = t.cell(0, i)
        shade(c, "EEF2FF" if i % 2 == 0 else "EDE9FE")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold, r.font.size, r.font.color.rgb = True, Pt(8), RGBColor.from_string(NAVY)
    doc.add_paragraph()


cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("How CV Upload Works in JobAI Scout")
r.font.name, r.font.size, r.font.bold, r.font.color.rgb = "Aptos Display", Pt(26), True, RGBColor.from_string(NAVY)
sub = doc.add_paragraph("A code-based explanation of private file upload, AI extraction, profile merge and security")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb, sub.runs[0].font.size = RGBColor.from_string(SLATE), Pt(11)
doc.add_paragraph()
info("Purpose:", "This document explains the existing CV Upload page and analyze-cv Edge Function. It is written for a reader who wants to understand what happens to a CV, how extracted information reaches the user profile, and where security checks belong.")

heading("1. The feature in one sentence")
para("A signed-in job seeker uploads a CV to a private Supabase Storage folder. The application stores the file path in the user profile, asks a protected server function to extract text and analyze it with AI, shows the result, and fills only empty profile fields so existing information is not overwritten.")
flow(["Select CV", "Upload to private storage", "Save file path", "Extract readable text", "AI returns structured data", "Fill empty profile fields", "Show merge result"])
grid(["Layer", "Responsibility", "Main code"], [
    ["User interface", "Lets the user choose or drag a file, shows upload/analysis status, extraction metadata and merge review.", "src/pages/CVUpload.tsx"],
    ["Private storage", "Stores the original resume under the authenticated user’s own storage folder.", "Supabase Storage bucket: resumes"],
    ["Analysis boundary", "Verifies the session, downloads the resume, extracts text and calls the AI provider.", "supabase/functions/analyze-cv/index.ts"],
    ["Extraction helpers", "Reads PDF/DOCX text and falls back to OCR when ordinary PDF text extraction is weak.", "supabase/functions/_shared/cv-extraction.ts"],
    ["Merge helpers", "Normalizes AI output and produces a safe profile update containing only missing values.", "cv-profile-merge.ts and src/lib/cv-extracted-data.ts"],
    ["Profile database", "Stores reusable career details such as skills, education, links and CV summary.", "public.profiles with RLS"],
], [3.2, 7.0, 6.4])

heading("2. What the user sees on the Upload CV page")
para("The page is protected for the job-seeker dashboard. It obtains the active user and current profile from AuthContext. Before upload, it calculates profile readiness from key fields such as name, email, phone, location, skills, desired roles, experience, resume path, LinkedIn, GitHub, portfolio, company and education. This gives the user a visible reason to upload a CV: the extracted information can improve job matching and application autofill.")
grid(["UI element", "How it works"], [
    ["Drop zone / file picker", "Accepts PDF, DOCX and DOC selection. Drag and drop explicitly rejects unsupported file types before upload."],
    ["Upload & Analyze button", "Starts only after a file and authenticated user are present. It changes state from uploading to analyzing."],
    ["Extraction information", "Displays extraction method, page count, character count and whether OCR was used when the Edge Function returns this metadata."],
    ["Merge Review", "Shows each field as extracted, preserved or skipped. Existing profile values are intentionally kept."],
    ["Extracted Data card", "Presents structured information such as skills, contact details, experience, education and summary."],
], [4.4, 12.2])
info("User-control principle:", "The feature is designed to enrich a profile, not replace it. If a profile field already has a meaningful value, the default merge logic preserves it rather than silently overwriting it with an AI interpretation.", "DCFCE7")

heading("3. Step-by-step upload flow")
grid(["Step", "What happens", "Data involved"], [
    ["1. File selection", "The browser gives CVUpload a File object after a user action. The page resets old extraction state so the new result cannot be confused with the previous file.", "File name, MIME type, local browser file object."],
    ["2. Create storage path", "The page builds a path: user.id / timestamp_fileName.", "The authenticated Supabase user ID becomes the first path segment."],
    ["3. Upload", "The browser calls supabase.storage.from('resumes').upload(filePath, file).", "The original document is stored in the private resumes bucket."],
    ["4. Save reference", "The page updates profiles.resume_url with the storage path, not a public download URL.", "Profile contains a private storage reference."],
    ["5. Invoke analysis", "The browser calls Supabase Function analyze-cv with fileName and filePath. The Supabase client sends the signed-in access token.", "Function request contains a file reference, not the password or a provider secret."],
    ["6. Show and merge", "Returned AI data is normalized, compared with a fresh profile query, and saved only where values are absent.", "Structured fields plus merge plan and filled field list."],
], [1.2, 9.0, 6.4])

heading("4. Why the resume is private")
para("The resumes storage bucket is created as non-public. The storage policies check the first folder name in the object path. A user may insert, read or update a file only when that folder equals auth.uid(), the ID from their verified Supabase access token. Because the app writes a path beginning with user.id, normal browser uploads follow this rule automatically.")
grid(["Policy", "Plain-language meaning"], [
    ["Users can upload resume", "A signed-in user may upload only to a path whose first folder is their own user ID."],
    ["Users can view own resume", "A signed-in user may download only documents stored under their own user ID folder."],
    ["Users can update own resume", "A signed-in user may update only a document stored under their own folder."],
    ["profiles RLS", "The user can update their own profile row, including resume_url, but another user cannot update it through the normal client."],
], [5.1, 11.5])
info("Important concept:", "The path stored in profiles.resume_url is a storage object key such as user-id/timestamp_resume.pdf. It is not automatically a public internet URL. A private file needs an authorized download or signed URL before it can be opened.")

heading("5. The protected analyze-cv Edge Function")
para("The browser does not send the entire CV directly to an AI API from frontend code. Instead, it invokes analyze-cv, a Supabase Edge Function. The function first requires an Authorization header. It creates an anonymous Supabase client to validate the bearer token and obtain the real authenticated user. Only then does it continue with the requested file and AI work.")
flow(["Browser invokes analyze-cv", "Function reads Authorization header", "Supabase verifies JWT", "Function gets authenticated user", "Function downloads CV", "Text extraction", "AI structured extraction", "Profile update + response"])
grid(["Function stage", "Implementation behavior"], [
    ["Authentication", "Rejects missing Authorization header. Calls anonClient.auth.getUser(token); rejects invalid or expired tokens."],
    ["Server credentials", "Uses SUPABASE_SERVICE_ROLE_KEY for protected server-side database/storage work. This key stays in server secrets, never in VITE frontend variables."],
    ["Text extraction", "Uses an optional Python extraction service first if configured; otherwise uses Deno helpers."],
    ["AI extraction", "Sends extracted plain text to OpenRouter with Gemini 2.5 Flash and requests a strict JSON object."],
    ["Profile persistence", "Builds an update from missing fields and uses upsert so an older account without a profile row can be repaired."],
    ["Response", "Returns normalized fields plus extraction metadata and a list of fields saved by the server."],
], [4.0, 12.6])

heading("6. How text is extracted from PDF and DOCX")
para("AI models work better when they receive clean text than when they receive arbitrary binary files. The shared extraction module therefore converts the uploaded document into bounded text before structured AI analysis. The system records metadata so the UI can explain what method was used.")
grid(["File situation", "Extraction behavior"], [
    ["Configured Python service available", "Sends the file to CV_EXTRACTOR_URL /extract. The service can use PDF and document libraries and returns text, pages, method and OCR status. It limits received files to 20 MB."],
    ["Normal text PDF", "Uses unpdf to extract per-page text and inserts page markers."],
    ["Scanned / image-heavy PDF", "If average extracted text is below 50 characters per page, the helper can call Gemini OCR using the protected server key."],
    ["DOCX", "Uses Mammoth to obtain raw text."],
    ["DOC or extraction failure", "Attempts a raw text decode as a last-resort fallback; output quality may be limited."],
    ["Long document", "Cleans blank lines and limits text retained for extraction; the AI request receives up to 15,000 characters of CV text."],
], [4.6, 12.0])
info("Accuracy note:", "OCR and AI extraction can make mistakes. The prompt tells the model not to invent details, but the user should still inspect the merge review and profile before relying on the result for applications.", "FEF3C7")

heading("7. What the AI is asked to return")
para("The Edge Function uses a low-temperature prompt and asks for JSON rather than free-form prose. It requests full name, email, phone, location, LinkedIn, GitHub, portfolio URL, current company, skills, suggested roles, total experience years, education, certifications, languages and a detailed CV summary. The prompt explicitly instructs the model to use only facts present in the CV.")
grid(["Raw AI property", "Profile destination"], [
    ["fullName, email, phone, location", "full_name, email, phone, location"],
    ["linkedinUrl, githubUrl, portfolioUrl", "linkedin_url, github_url, portfolio_url"],
    ["currentCompany, experienceYears", "current_company, experience_years"],
    ["skills, suggestedRoles", "skills, desired_roles"],
    ["education, certifications, languages", "education, certifications, languages"],
    ["cvSummary", "cv_summary and bio if bio is also empty"],
], [6.2, 10.4])

heading("8. Normalization and safe profile merge")
para("Raw AI JSON is not trusted as-is. The normalizeExtractedData helper accepts both camelCase and snake_case key variants, trims strings, converts comma-separated strings into arrays when necessary, and accepts positive numeric experience values. After normalization, buildProfileUpdateFromExtracted compares every possible extracted field against the current profile.")
grid(["Merge rule", "Result"], [
    ["Existing profile value is empty and extracted value is usable", "The field is added to updatePayload and recorded in filledKeys."],
    ["Existing profile value already has content", "The field is preserved. AI does not overwrite it by default."],
    ["AI returns blank / unusable value", "The field is skipped."],
    ["bio is empty but cvSummary is usable", "The same summary can fill both cv_summary and bio."],
    ["Full update cannot persist", "The client or function can fall back to a core set of profile fields and reports whether saving succeeded."],
], [6.1, 10.5])
para("The system also calls update_profile_data_sources when available. This lets the profile indicate that particular fields came from AI/CV extraction, which is useful for transparency and later editing decisions.")

heading("9. How the feature connects to later workflows")
grid(["Later feature", "How CV data helps"], [
    ["Profile readiness", "The stored resume path and extracted profile fields raise the completeness indicator."],
    ["Job matching", "Skills, desired roles, experience, location and summary give later matching features more structured input."],
    ["Application autofill", "The browser extension can reuse approved profile information instead of repeatedly asking for the same details."],
    ["Cover letters / career AI", "Structured profile data and CV context can guide personalized output, subject to the applicable Edge Function policies."],
    ["Recruiter workflow", "Recruiters can review applicant profile information only through the intended application-linked access policies."],
], [5.0, 11.6])

heading("10. Current security and reliability improvements to make")
para("The existing design has good foundations: authenticated user folders, a non-public resume bucket, token validation, protected secrets and conservative merge behavior. The following improvements are important before claiming full production readiness.")
grid(["Finding", "Why it matters", "Recommended improvement"], [
    ["Service-role download path is not checked against the caller", "Service-role access bypasses Storage RLS. An authenticated caller who guesses another file path could ask the function to download it.", "Before download, require filePath to start with `${user.id}/`; reject every other path. Prefer a user-scoped storage client when possible."],
    ["No clear size/type enforcement inside Edge Function", "A browser can bypass the UI accept attribute; large or unexpected uploads can increase cost or fail unpredictably.", "Validate extension, MIME type and maximum file size in the client and again in the Edge Function before processing."],
    ["Old .doc support is weak", "The Deno fallback can only raw-decode legacy .doc, so results may be poor.", "Officially support PDF and DOCX; convert legacy DOC server-side or reject it with clear guidance."],
    ["AI extraction is probabilistic", "Even a constrained model may misread dates, skills or OCR text.", "Add explicit per-field user approval/edit controls before persistence for high-stakes data."],
    ["CORS allows all origins", "Broad CORS is often unnecessary for a single production frontend.", "Restrict Access-Control-Allow-Origin to approved production and local development origins where compatible."],
], [4.2, 6.1, 6.3])
info("Do not expose secrets:", "OPENROUTER_API_KEY, SUPABASE_SERVICE_ROLE_KEY and CV_EXTRACTOR_URL configuration belong to protected server configuration. They must never be placed in Vercel VITE variables, browser code, screenshots, or Git commits.", "FEE2E2")

heading("11. Reader summary")
para("In short, JobAI Scout stores the original CV privately under the user’s own storage folder, transforms the file into readable text on the server, asks an AI model for a controlled JSON summary, and fills only missing career-profile fields. The privacy model depends on Storage RLS for normal uploads and on strong ownership validation in every service-role function. The main hardening task is to verify the caller owns filePath before the analysis function downloads any resume using its privileged server credential.")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("JobAI Scout CV Upload - Technical Explanation")
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor.from_string(SLATE)

doc.core_properties.title = "How CV Upload Works in JobAI Scout"
doc.core_properties.author = "JobAI Scout"
doc.save(OUT)
print(OUT)
