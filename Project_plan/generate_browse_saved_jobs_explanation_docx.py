"""Generate a reader-friendly explanation of the Browse Jobs and Saved Jobs features."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name("JobAI_Scout_Browse_and_Saved_Jobs_Explained.docx")
NAVY, INDIGO, VIOLET, SLATE = "0B1028", "4F46E5", "7C3AED", "334155"


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), color)
    props.append(node)


def bottom_border(paragraph):
    props = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), INDIGO)
    border.append(bottom)
    props.append(border)


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Cm(1.7)
section.left_margin = section.right_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Aptos", Pt(10)
normal.font.color.rgb = RGBColor.from_string("172033")
normal.paragraph_format.space_after = Pt(6)
for name, size, color in [("Heading 1", 17, NAVY), ("Heading 2", 13, INDIGO), ("Heading 3", 11, VIOLET)]:
    style = doc.styles[name]
    style.font.name, style.font.size, style.font.bold = "Aptos Display", Pt(size), True
    style.font.color.rgb = RGBColor.from_string(color)


def heading(text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        bottom_border(p)
    return p


def para(text):
    return doc.add_paragraph(text)


def info(label, text, color="EEF2FF"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, color)
    p = cell.paragraphs[0]
    p.add_run(label + " ").bold = True
    p.add_run(text)
    doc.add_paragraph()


def grid(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        shade(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.bold, run.font.color.rgb = True, RGBColor(255, 255, 255)
    for row_number, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_number % 2 == 0:
                shade(cells[i], "F8FAFC")
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()


def flow(steps):
    table = doc.add_table(rows=1, cols=len(steps))
    for i, text in enumerate(steps):
        cell = table.cell(0, i)
        shade(cell, "EEF2FF" if i % 2 == 0 else "EDE9FE")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold, run.font.size, run.font.color.rgb = True, Pt(8), RGBColor.from_string(NAVY)
    doc.add_paragraph()


cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run("How Browse Jobs and Saved Jobs Work in JobAI Scout")
run.font.name, run.font.size, run.font.bold = "Aptos Display", Pt(25), True
run.font.color.rgb = RGBColor.from_string(NAVY)
sub = doc.add_paragraph("A clear explanation of live job collection, search, matching, bookmarks and privacy")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb, sub.runs[0].font.size = RGBColor.from_string(SLATE), Pt(11)
doc.add_paragraph()
info("Purpose:", "This document explains the existing Browse Jobs and Saved Jobs tabs. It describes what the user sees, how roles are collected and searched, how job bookmarks are stored, and how the system keeps each user's saved list private.")

heading("1. The feature in one sentence")
para("Browse Jobs gives a signed-in job seeker a searchable stream of active roles collected from selected sources. Saved Jobs lets that same user bookmark interesting roles and return to them later without saving anything for other users.")
flow(["Choose source", "Search and filter", "Collect or load roles", "Open a role", "Save a role", "Review Saved Jobs", "Apply when ready"])
grid(["Layer", "Responsibility", "Main code"], [
    ["Browse Jobs page", "Displays job discovery controls, search results, filters, saving, application links and tailored cover-letter action.", "src/pages/JobBoard.tsx"],
    ["Saved Jobs page", "Loads the signed-in user's bookmarks, separates extension recommendations from job-board roles, and lets the user remove a bookmark.", "src/pages/SavedJobs.tsx"],
    ["Collection service", "Authenticates the caller, invokes the chosen source adapter, filters irrelevant results, removes duplicates and stores valid jobs.", "supabase/functions/collect-jobs/index.ts"],
    ["Search function", "Filters active jobs and calculates a consistent keyword match score for collected roles.", "public.search_collected_jobs(...)"],
    ["Database and policies", "Stores jobs and bookmarks with Row Level Security, so a user sees and removes only their own bookmarks.", "public.jobs and public.saved_jobs"],
], [3.2, 7.5, 5.9])

heading("2. Browse Jobs: what the user sees")
para("The Browse Jobs tab is the main discovery workspace. Its header explains the job-discovery purpose, shows the number of indexed roles, and asks the user to choose one source before refreshing. This one-source-at-a-time design avoids a slow provider blocking every other source and gives clearer feedback when a source has a problem.")
grid(["Interface element", "What it does"], [
    ["Source buttons", "Lets the user select LinkedIn, Indeed, RSS feeds, company career pages, or another configured source before collecting jobs."],
    ["Search input", "Accepts a job title, skill, or company. A refresh requires a search term, so requests sent to job providers remain useful and targeted."],
    ["Location input", "Narrows displayed roles and is also sent as collection context. The UI uses Pakistan as the default country context when a city is entered alone."],
    ["Filters", "Offers job type, source, remote/hybrid work mode, and a clear-filters action. Results update after a short debounce rather than on every keystroke."],
    ["Open-role cards", "Show source, company, title, location, job type, salary when available, a short description, skills, posting recency, save control, tailor-letter control and apply action."],
    ["Pagination", "Shows up to 30 collected roles per page, avoiding an overly heavy page when the job database grows."],
], [4.2, 12.4])
info("Important user experience decision:", "The page does not claim that every job is a perfect match. It labels source, freshness and match context clearly, then leaves the final decision to the job seeker.", "DCFCE7")

heading("3. Browse Jobs: step-by-step data flow")
grid(["Step", "What happens", "Privacy / reliability detail"], [
    ["1. Authentication", "JobBoard reads the current user from AuthContext. If no user is present, collection and saved-job actions stop.", "The browser uses the signed-in Supabase session; no provider secret is present in the UI."],
    ["2. Load roles", "The page calls search_collected_jobs with normalized search terms, source, location, job type, work mode, page limit and offset.", "Terms are cleaned, limited to four and must contain at least two characters."],
    ["3. Choose a source", "The user selects one provider group and enters a search term before pressing Refresh.", "The UI blocks an all-sources refresh, making source limits and failures easier to manage."],
    ["4. Collect", "The browser invokes collect-jobs. The Edge Function validates the session, selects the correct adapter, collects roles, filters them and upserts them.", "Provider keys and the service role key remain server-side in Supabase secrets."],
    ["5. Display", "The page reloads page one and presents active roles with direct posting links. Realtime listening also refreshes when recommended-job data changes.", "Only jobs with a direct source URL or a recruiter owner are displayed in the modern open-role list."],
    ["6. Act", "The user can bookmark, tailor a cover letter, or apply. External application links open in a separate protected tab.", "Applying through an external URL is intentionally not treated as proof that the third-party application was submitted."],
], [1.2, 8.7, 6.7])

heading("4. How job collection works")
para("The collection function is a server-side boundary between the browser and job providers. It requires an Authorization header, verifies the Supabase user, and then chooses only the requested source. Depending on the source, it can call an Indeed adapter, LinkedIn adapter, multi-job-board adapter, configured RSS feed, or configured company-career-page adapter.")
flow(["User refreshes one source", "Function validates session", "Adapter fetches roles", "Relevance filter", "Duplicate removal", "Upsert into jobs", "Return useful result summary"])
grid(["Collection safeguard", "Why it exists"], [
    ["One source per refresh", "Prevents one provider from making every refresh slow or unreliable and respects provider-specific limits."],
    ["Maximum item limit", "Limits a request to a small collection budget, which controls response time and external API cost."],
    ["Source status updates", "Configured RSS and company sources record last collection time, result count and errors so an administrator can identify unhealthy sources."],
    ["Relevance check", "RSS and company sites can expose many unrelated jobs, so these sources need a threshold before roles are stored or shown."],
    ["Deduplication", "The collection helper removes repeated roles before the database upsert, reducing duplicate job cards."],
    ["Direct-link requirement", "The modern listing hides roles that lack both a direct source URL and a recruiter owner, avoiding frustrating cards that cannot be acted on."],
], [5.0, 11.6])

heading("5. Search, filtering and match score")
para("search_collected_jobs is a database function that works over active jobs. It normalizes the search terms, then scores title coverage at 50%, skills coverage at 30%, and description coverage at 20%. For strict RSS and company-career searches, the role must score at least 30% and match at least one term in the title or skills. This keeps broad feeds useful without making a score look like a prediction of hiring success.")
grid(["Filter or score input", "Implementation behavior"], [
    ["Search terms", "Lowercased, trimmed, deduplicated and limited to four terms. Punctuation is removed except useful characters such as +, # and -."],
    ["Source", "Matches the stored source name, such as linkedin_apify, indeed_apify, rss or company_career."],
    ["Location", "Performs a case-insensitive partial match against the job location."],
    ["Job type / work mode", "Filters stored job type and remote/hybrid context."],
    ["Order", "Sorts by calculated score, then most recent posting date, then creation date."],
    ["Page size", "Caps each request at 30 roles and safely bounds the offset."],
], [4.4, 12.2])
info("What the score means:", "It is a keyword relevance score for search and discovery. It is not an AI judgment about a candidate's worth, qualification, or chance of being hired.", "FEF3C7")

heading("6. Saving a job")
para("A bookmark is a private record in public.saved_jobs. On the Browse Jobs page, selecting the bookmark icon inserts a row containing the current user ID and job ID. Selecting it again deletes that row. The interface updates its local saved-ID set immediately, so the bookmark icon changes without requiring a full page reload.")
grid(["Database field", "Purpose"], [
    ["id", "Unique identifier for the bookmark row."],
    ["user_id", "The authenticated account that owns the bookmark."],
    ["job_id", "The job-board role being saved. Legacy support also allows recommended_job_id for extension-generated recommendations."],
    ["saved_at", "Timestamp used to show the most recently saved jobs first."],
    ["Unique user/job rule", "Prevents the same user saving the same job repeatedly."],
], [4.4, 12.2])
flow(["Click bookmark", "Insert saved_jobs row", "RLS checks user ID", "Update icon state", "Saved Jobs loads it later"])

heading("7. Saved Jobs tab: what it does")
para("The Saved Jobs tab is a personal shortlist. It fetches bookmarks for the signed-in user and displays two groups so the origin is transparent: roles saved from extension scans and regular roles saved from the job board. Both sections display useful context such as company, location, salary or skills when those details exist.")
grid(["Saved Jobs behavior", "How it works"], [
    ["Load bookmarks", "Queries saved_jobs where user_id equals the active user, ordered by saved_at descending."],
    ["Extension scan group", "Loads recommended_jobs referenced by recommended_job_id, then reconstructs the saved order using a map."],
    ["Job board group", "Uses the jobs(*) relationship on saved_jobs rows whose recommended_job_id is null."],
    ["Apply action", "Uses source_url for extension recommendations and job_url for older regular jobs, opening the original job page in a new tab."],
    ["Remove action", "Deletes the saved_jobs row by its own ID and removes the card from the current interface state."],
    ["Empty state", "Clearly tells a user there are no bookmarks and directs them to browse and save roles."],
], [4.4, 12.2])

heading("8. Data protection and access control")
para("The saved-jobs system relies on Supabase Row Level Security (RLS). The browser uses the anonymous publishable key and the user's session. Database policies compare auth.uid(), derived from the verified access token, with the saved_jobs.user_id column. Therefore a normal user cannot list, create, or delete another user's saved roles through the application API.")
grid(["Policy", "Plain-language meaning"], [
    ["Users can view their saved jobs", "A user can select only bookmarks where the bookmark belongs to their own authenticated account."],
    ["Users can save jobs", "A user can insert a bookmark only when user_id equals their own authenticated account."],
    ["Users can unsave jobs", "A user can delete only their own bookmark rows."],
    ["Jobs view policy", "Authenticated users may read active job records needed to discover and open roles."],
    ["Server collection secrets", "Provider tokens and the Supabase service role key are used only in Edge Functions, never sent to the browser."],
], [4.8, 11.8])
info("Privacy boundary:", "Saving a job is not an application. It is a private bookmark. Recruiters and job sources do not receive a notification merely because a user saved their role.", "DCFCE7")

heading("9. Related actions on a job card")
grid(["Action", "What it does", "What it does not do"], [
    ["Save", "Stores a private bookmark for later review.", "Does not apply or share the user's data."],
    ["Tailor letter", "Securely asks the server to use the stored job and signed-in profile to create an editable, professional cover letter.", "Does not automatically submit an application."],
    ["Apply", "Opens the employer, job board, or recruiter application destination.", "Does not guarantee the external site accepted the application."],
    ["Remove saved", "Deletes only the user's bookmark row.", "Does not delete the underlying job or affect another user's saved list."],
], [3.3, 7.1, 6.0])

heading("10. Current reliability and future improvements")
grid(["Area", "Current design", "Useful next improvement"], [
    ["Provider availability", "The UI shows source-specific errors and collection keeps fulfilled source results even when another task fails.", "Add a visible source-health panel using the stored last_error and last_collected_at values."],
    ["Saved job history", "Bookmarks remain until the user removes them; the page says they are never auto-deleted.", "Show a subtle warning if the original posting URL is unavailable or the job has become inactive."],
    ["Application tracking", "Internal recruiter applications can create application records; external links open separately.", "Let users manually mark an external application as applied after completing it."],
    ["Search relevance", "Uses transparent keyword scoring and a stricter threshold for broad feeds.", "Offer optional profile-aware ranking while keeping the reason for each score visible."],
    ["Network feedback", "Loading, success and error toasts exist for primary actions.", "Add retry controls for individual failed source collections without re-running a full search."],
], [3.4, 6.3, 6.7])

heading("11. Reader summary")
para("JobAI Scout's Browse Jobs tab collects and searches active jobs through authenticated server-side adapters, then presents clean, paginated role cards with clear source and relevance context. Saved Jobs is a private, user-owned shortlist built on saved_jobs and Row Level Security. The user remains in control at every stage: they choose a source, search and filter roles, bookmark only what matters, review an AI-tailored letter before use, and decide when to apply on the original job site.")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("JobAI Scout Browse Jobs & Saved Jobs - Technical Explanation")
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor.from_string(SLATE)

doc.core_properties.title = "How Browse Jobs and Saved Jobs Work in JobAI Scout"
doc.core_properties.author = "JobAI Scout"
doc.save(OUT)
print(OUT)
