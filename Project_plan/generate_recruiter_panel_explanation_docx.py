"""Generate the complete JobAI Scout recruiter panel explanation."""

from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name("JobAI_Scout_Recruiter_Panel_Explained.docx")
NAVY, INDIGO, VIOLET, SLATE = "0B1028", "4F46E5", "7C3AED", "334155"


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), color)
    props.append(node)


def bottom_border(paragraph):
    props = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), INDIGO)
    border.append(bottom)
    props.append(border)


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Cm(1.65)
section.left_margin = section.right_margin = Cm(1.8)
normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Aptos", Pt(10)
normal.font.color.rgb = RGBColor.from_string("172033")
normal.paragraph_format.space_after = Pt(6)
for name, size, color in [("Heading 1", 17, NAVY), ("Heading 2", 13, INDIGO), ("Heading 3", 11, VIOLET)]:
    style = doc.styles[name]
    style.font.name, style.font.size, style.font.bold = "Aptos Display", Pt(size), True
    style.font.color.rgb = RGBColor.from_string(color)


def heading(text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        bottom_border(p)
    return p


def para(text):
    return doc.add_paragraph(text)


def info(label, text, color="EEF2FF"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, color)
    p = cell.paragraphs[0]
    p.add_run(label + " ").bold = True
    p.add_run(text)
    doc.add_paragraph()


def grid(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        shade(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row_no, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_no % 2 == 0:
                shade(cells[i], "F8FAFC")
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()


def flow(steps):
    table = doc.add_table(rows=1, cols=len(steps))
    for i, text in enumerate(steps):
        cell = table.cell(0, i)
        shade(cell, "EEF2FF" if i % 2 == 0 else "EDE9FE")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold, run.font.size = True, Pt(8)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph()


cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
title = cover.add_run("How the JobAI Scout Recruiter Panel Works")
title.font.name, title.font.size, title.font.bold = "Aptos Display", Pt(25), True
title.font.color.rgb = RGBColor.from_string(NAVY)
subtitle = doc.add_paragraph("Complete recruitment workflow, data-flow, privacy and technical explanation")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.color.rgb, subtitle.runs[0].font.size = RGBColor.from_string(SLATE), Pt(11)
doc.add_paragraph()
info("Purpose:", "This document explains the existing JobAI Scout recruiter workspace in reader-friendly language. It covers recruiter authorization, the company profile, job posting and editing, applicant access, pipeline decisions, private notes, application status, database protection, user experience, current limitations and a recommended operating workflow.")

heading("1. The recruiter panel in one sentence")
para("The recruiter panel is a protected hiring workspace where a verified recruiter represents a company, publishes and maintains its own jobs, sees only candidates who applied to those jobs, organizes them through a simple pipeline, and records private review notes.")
flow(["Recruiter signs in", "Complete company profile", "Post a job", "Candidate applies", "Review candidate", "Update status", "Track outcome"])

heading("2. Recruiter modules and routes")
grid(["Module", "Route", "Responsibility"], [
    ["Company Profile", "/recruiter/profile", "Stores the company identity reused when the recruiter creates jobs."],
    ["Post a Job", "/recruiter/jobs?new=1", "Opens the job dialog immediately through a navigation query parameter."],
    ["My Jobs", "/recruiter/jobs", "Lists, creates, edits and deletes only the signed-in recruiter's listings."],
    ["Applicants", "/recruiter/candidates", "Shows the candidate pipeline, limited profile evidence, status controls and private notes."],
    ["Application Status", "/recruiter/application-status", "Provides a simplified read-only summary of application outcomes by job."],
], [3.4, 4.0, 9.2])

heading("3. Authentication and recruiter authorization")
para("Supabase authenticates the account first. AuthContext then loads the profile and role from profiles and user_roles. If the role is recruiter, it also loads recruiter_profiles. Every recruiter route uses ProtectedRoute with requiredRole=\"recruiter\".")
flow(["Supabase session", "Load role", "Load recruiter profile", "ProtectedRoute verifies recruiter", "Render or redirect"])
grid(["Account condition", "Result"], [
    ["Not signed in", "Redirect to /login."],
    ["Job-seeker role", "Redirect to /dashboard."],
    ["Admin role", "Redirect to /admin."],
    ["Recruiter role", "Allow the requested recruiter route."],
    ["Opening /recruiter", "Redirect to /recruiter/jobs."],
], [5.3, 11.3])
info("Defense in depth:", "The route guard controls the interface, while Supabase Row Level Security controls database rows. A recruiter should never gain another recruiter's jobs or candidates merely by changing a URL or request in the browser.", "DCFCE7")

heading("4. Shared recruiter interface")
grid(["Interface area", "How it helps"], [
    ["Recruiter sidebar", "Shows Company Profile, Post a Job, My Jobs, Applicants and Application Status."],
    ["Role identity", "The shared dashboard shell labels the workspace Recruiter Portal and uses a cyan recruiter accent."],
    ["Post shortcut", "The Post a Job item uses ?new=1; RecruiterJobs reads it, opens the dialog and removes the query parameter."],
    ["Responsive layout", "Cards, forms and applicant details collapse into mobile-friendly columns while preserving desktop density."],
    ["Status feedback", "Loading states and toasts explain database progress, success and errors."],
    ["Consistent brand", "The JobAI Scout signal mark and dark navy visual system are shared with the public site and other workspaces."],
], [4.2, 12.4])

heading("5. Company Profile")
para("The company profile is the recruiter's reusable organization identity. It prevents repeated typing and gives candidates consistent employer information across new postings.")
grid(["Field", "Use"], [
    ["Company name", "Required. Used as the fallback company name when a new job does not specify one."],
    ["Website", "Optional organization website."],
    ["Industry", "Optional business category such as Technology or Finance."],
    ["Description", "Optional candidate-facing summary of the organization."],
], [4.0, 12.6])
flow(["Load recruiterProfile", "Populate form", "Validate company name", "Trim optional values", "Upsert by user_id", "Show result"])
info("Ownership:", "The saved row includes the current authenticated user_id. Database policies allow recruiters to insert, view and update their own recruiter profile; administrators may inspect or update profiles under admin policy.")

heading("6. Creating a job posting")
grid(["Job field", "Saved behavior"], [
    ["Title", "Required by the UI before Post Job becomes available."],
    ["Company", "Uses the typed company or falls back to recruiterProfile.company_name."],
    ["Location", "Optional; blank becomes null."],
    ["Employment type", "Defaults to full-time."],
    ["Experience level", "Optional entry, mid, senior or lead value."],
    ["Salary minimum/maximum", "Optional text input converted to a number before storage."],
    ["Description", "Optional detailed role explanation."],
    ["Skills", "Comma-separated text normalized into a trimmed array."],
    ["Requirements", "Comma-separated text normalized into a trimmed array."],
    ["Application URL", "Optional destination for an external application flow."],
    ["System ownership", "recruiter_id is the active user and source is always recruiter."],
], [4.5, 12.1])
flow(["Open dialog", "Enter job facts", "Normalize values", "Attach recruiter_id", "Insert jobs row", "Refresh My Jobs"])

heading("7. Editing and deleting jobs")
grid(["Action", "Current behavior"], [
    ["Edit", "Copies an existing job into the same dialog, records editId, and updates that row when saved."],
    ["After save", "Closes the dialog, clears form/edit state and reloads jobs newest first."],
    ["Delete", "Deletes the selected jobs row and reloads the recruiter's list."],
    ["Ownership boundary", "Database policies permit update and deletion only where recruiter_id equals the authenticated user and the caller has recruiter role."],
], [4.1, 12.5])
info("Destructive-action note:", "Job deletion currently happens immediately without a confirmation dialog. A competition-grade version should name the job, explain applicant impact and require explicit confirmation.", "FEE2E2")

heading("8. How a candidate enters the pipeline")
flow(["Candidate browses active job", "Clicks Apply", "Insert job_applications", "Status defaults to New", "Recruiter's RLS permits view", "Applicant appears in pipeline"])
para("The candidate application contains user_id, job_id, applied_at and status. A database migration normalizes old null/applied values to new, makes new the default, and restricts the pipeline to New, Shortlisted, Rejected or Hired.")
grid(["Status", "Meaning"], [
    ["New", "The application has arrived and has not yet been advanced."],
    ["Shortlisted", "The recruiter wants to continue evaluation or interviewing."],
    ["Rejected", "The candidate is no longer progressing for that application."],
    ["Hired", "The application reached the successful outcome."],
], [4.0, 12.6])

heading("9. Applicant Pipeline")
para("RecruiterCandidates loads applications joined to job title, company and recruiter ownership, keeps the recruiter's owned applications, then loads only the corresponding applicant profiles and private notes.")
grid(["Visible information", "Purpose and boundary"], [
    ["Candidate identity", "Full name and email support communication and identification."],
    ["Application context", "Job title, company and application date keep review tied to the correct role."],
    ["Experience and location", "Provides high-level screening context without exposing the entire account."],
    ["Skills", "Displays up to eight saved skills for quick evidence-based review."],
    ["Resume indicator", "Shows whether resume_url exists; the current panel does not provide a resume download/view button."],
    ["Pipeline counters", "Calculates totals for New, Shortlisted, Rejected and Hired from the loaded applications."],
], [4.4, 12.2])

heading("10. Candidate privacy and database boundaries")
grid(["Rule", "Protection"], [
    ["Application ownership", "A recruiter can read an application only when its job belongs to that recruiter."],
    ["Candidate profile access", "A recruiter can read a profile only when that candidate applied to one of the recruiter's jobs."],
    ["Status updates", "A recruiter may update applications attached to owned jobs; the database status constraint prevents unsupported states."],
    ["Job ownership", "Insert, update and delete rules require recruiter role plus matching recruiter_id."],
    ["Private notes", "A recruiter manages notes where recruiter_id equals their authenticated ID."],
    ["Candidate note visibility", "Candidates may read only public notes about themselves; this UI deliberately creates is_private=true notes."],
], [4.5, 12.1])
info("Privacy principle:", "Recruiter access is created by a real application relationship, not by the recruiter role alone. This is much safer than allowing every recruiter to browse every user's profile.", "DCFCE7")

heading("11. Updating application status")
flow(["Choose status", "Disable application control", "Update job_applications", "RLS checks owned job", "Update local card", "Show confirmation"])
para("The status selector is labeled for accessibility using the candidate's name. While a change is saving, that application's control is disabled. A successful response updates local state without requiring a full page reload; a failed response leaves the old status and shows a destructive toast.")

heading("12. Private recruiter notes")
flow(["Expand candidate", "Write review note", "Insert candidate_notes", "Set is_private=true", "Group by candidate + job", "Display newest notes"])
grid(["Design choice", "Reason"], [
    ["Candidate plus job key", "The same person can apply to multiple roles without mixing evaluation notes."],
    ["Recruiter ID saved", "Ownership is explicit and enforceable by RLS."],
    ["Private flag always true", "The applicant cannot see interview or internal evaluation notes through this feature."],
    ["Empty note blocked", "The button remains disabled until trimmed text exists."],
    ["Immediate local update", "A saved note appears at the top without reloading the entire applicant list."],
], [4.4, 12.2])

heading("13. Application Status summary")
para("RecruiterApplicationStatus is a simpler, read-only operational view. It loads applications for jobs owned by the recruiter, calculates the same four status totals, and lists the job, company, application date and current status badge.")
grid(["Applicants page", "Application Status page"], [
    ["Interactive pipeline with candidate details", "Compact overview grouped visually by status."],
    ["Can change application status", "Read-only; reflects changes made elsewhere."],
    ["Shows skills, experience, location and resume indicator", "Shows job, company, applied date and status."],
    ["Supports private notes", "Does not display or create notes."],
], [8.3, 8.3])

heading("14. Loading, empty and error states")
grid(["State", "Recruiter experience"], [
    ["No company name", "Save is stopped and a clear Company name is required toast appears."],
    ["No jobs", "My Jobs explains that the recruiter can begin with Post Job."],
    ["No applicants", "The pipeline explains that applications will appear after users apply."],
    ["No status activity", "Application Status shows a dedicated empty card."],
    ["Data loading", "Applicant screens show a centered loader while related records are assembled."],
    ["Database failure", "Applicant, profile, note and status operations surface readable destructive toasts."],
    ["Save in progress", "Buttons or application controls disable to reduce repeated writes."],
], [4.3, 12.3])

heading("15. Current implementation notes and improvement priorities")
grid(["Current observation", "Recommended competition-grade improvement"], [
    ["Job deletion has no confirmation and ignores returned delete errors.", "Add an ownership-aware confirmation dialog, inspect the error and explain whether existing applications block or cascade."],
    ["Job save closes the dialog even when Supabase returns an error.", "Keep the form open on failure, preserve entered data and focus the problem summary."],
    ["Company profile context is not refreshed immediately after upsert.", "Refresh recruiterProfile after save so new job forms use the latest company without a page reload."],
    ["Resume is shown only as an attached indicator.", "Provide a protected signed download/view action after verifying recruiter-to-application ownership."],
    ["Applicants are queried broadly and then filtered again in the browser.", "Query by owned job IDs or use a recruiter pipeline RPC for efficiency; retain RLS as the final security boundary."],
    ["Applicants and Application Status overlap substantially.", "Differentiate them with filters, interviews, activity history and funnel analytics, or consolidate them."],
    ["No recruiter analytics route is currently exposed.", "Add recruiter-only hiring analytics for views, applications, conversion, time-to-review and source effectiveness."],
    ["No search, role filter, pagination or sorting controls exist in the pipeline.", "Add server-side filtering and pagination for realistic employer volumes."],
    ["No audit trail records who changed a candidate status and when.", "Add application-status history for accountability and team collaboration."],
], [7.2, 9.4])

heading("16. Recommended recruiter workflow")
grid(["Step", "Recruiter task", "Why it matters"], [
    ["1", "Complete and verify the Company Profile.", "Keeps every new listing credible and consistently branded."],
    ["2", "Post a complete job with realistic skills, requirements, salary and application destination.", "Improves candidate understanding and matching quality."],
    ["3", "Review My Jobs after publishing and correct inaccurate details.", "Prevents candidates applying against outdated requirements."],
    ["4", "Open Applicants regularly and review New candidates using saved evidence.", "Reduces response delays and keeps decisions consistent."],
    ["5", "Use Shortlisted only for candidates who genuinely advance.", "Maintains a trustworthy pipeline rather than a loose bookmark list."],
    ["6", "Record concise job-specific private notes without sensitive or discriminatory content.", "Supports accountable evaluation while protecting candidate dignity."],
    ["7", "Use Application Status to verify overall outcomes and stalled applications.", "Provides a calm operational check after detailed reviews."],
], [1.0, 8.0, 7.6])

heading("17. Reader summary")
para("JobAI Scout's recruiter panel connects company identity, owned job listings and applicant evidence through a clear role-protected workflow. Recruiters can create and maintain their own jobs, see candidates only after a valid application relationship exists, move applications through four predictable stages and store recruiter-owned private notes. React provides the responsive interface and feedback, Supabase stores the data, and Row Level Security enforces ownership. The current system is a strong recruitment foundation; confirmation flows, protected resume access, recruiter analytics, scalable filtering and an audit history are the most valuable next improvements.")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("JobAI Scout Recruiter Panel - Complete Technical Explanation")
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor.from_string(SLATE)
doc.core_properties.title = "How the JobAI Scout Recruiter Panel Works"
doc.core_properties.author = "JobAI Scout"
doc.core_properties.subject = "Recruiter profile, job posting, applicant pipeline, privacy and hiring workflow"
doc.save(OUT)
print(OUT)
