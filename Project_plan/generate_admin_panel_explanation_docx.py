"""Generate a reader-friendly technical explanation of the JobAI Scout admin panel."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name("JobAI_Scout_Admin_Panel_Explained.docx")
NAVY, INDIGO, VIOLET, SLATE = "0B1028", "4F46E5", "7C3AED", "334155"


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), color)
    props.append(node)


def bottom_border(paragraph):
    props = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), INDIGO)
    border.append(bottom)
    props.append(border)


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Cm(1.65)
section.left_margin = section.right_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Aptos", Pt(10)
normal.font.color.rgb = RGBColor.from_string("172033")
normal.paragraph_format.space_after = Pt(6)
for name, size, color in [("Heading 1", 17, NAVY), ("Heading 2", 13, INDIGO), ("Heading 3", 11, VIOLET)]:
    style = doc.styles[name]
    style.font.name, style.font.size, style.font.bold = "Aptos Display", Pt(size), True
    style.font.color.rgb = RGBColor.from_string(color)


def heading(text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        bottom_border(paragraph)
    return paragraph


def para(text):
    return doc.add_paragraph(text)


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def info(label, text, color="EEF2FF"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, color)
    paragraph = cell.paragraphs[0]
    paragraph.add_run(label + " ").bold = True
    paragraph.add_run(text)
    doc.add_paragraph()


def grid(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        shade(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row_number, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = str(text)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_number % 2 == 0:
                shade(cells[index], "F8FAFC")
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()


def flow(steps):
    table = doc.add_table(rows=1, cols=len(steps))
    for index, text in enumerate(steps):
        cell = table.cell(0, index)
        shade(cell, "EEF2FF" if index % 2 == 0 else "EDE9FE")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold, run.font.size = True, Pt(8)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph()


cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
title = cover.add_run("How the JobAI Scout Admin Panel Works")
title.font.name, title.font.size, title.font.bold = "Aptos Display", Pt(25), True
title.font.color.rgb = RGBColor.from_string(NAVY)
subtitle = doc.add_paragraph("Complete interface, authorization, data-flow and operations explanation")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.color.rgb, subtitle.runs[0].font.size = RGBColor.from_string(SLATE), Pt(11)
doc.add_paragraph()
info("Purpose:", "This document explains the real JobAI Scout admin implementation in clear language. It covers access control, dashboard statistics, user and role management, job-source administration, platform analytics, the voice/RAG control room, backend communication, safety boundaries, current limitations and a practical operating workflow.")

heading("1. The admin panel in one sentence")
para("The JobAI Scout admin panel is a protected operations workspace where an authorized administrator can understand platform activity, manage users and roles, control job data sources and listings, review analytics, and inspect or configure the AI voice knowledge pipeline.")
flow(["Admin signs in", "Role is loaded", "Route is authorized", "Panel reads Supabase", "Admin reviews evidence", "Action reaches backend", "UI refreshes"])

heading("2. Admin modules and routes")
grid(["Module", "Route", "What it provides"], [
    ["Platform overview", "/admin", "Counts users, jobs and applications, then summarizes extension fill activity."],
    ["Manage Users", "/admin/users", "Lists profiles, roles and application counts; edits profiles and promotes or demotes administrators."],
    ["Manage Jobs", "/admin/jobs", "Manages collection sources and reviews or deletes stored job listings."],
    ["Platform Analytics", "/admin/analytics", "Reuses the main analytics engine to visualize active-job market demand and trends."],
    ["Voice Control Room", "/admin/voice", "Inspects voice usage, RAG search logs, indexed sources, retrieval matches and voice settings."],
], [3.4, 3.4, 9.8])
info("Navigation note:", "The first four areas are protected routes, but the current admin sidebar displays Overview, Users, Jobs and Platform Analytics only. The Voice Control Room route exists and works at /admin/voice, but a sidebar item has not yet been added.", "FEF3C7")

heading("3. Authentication and authorization")
para("Admin access begins with ordinary Supabase authentication, but authentication alone is not enough. AuthContext loads the signed-in account and reads its role from user_roles. ProtectedRoute compares that role with the admin requirement on every admin route.")
flow(["Supabase session", "Read profile + user_roles", "Resolve role", "ProtectedRoute checks admin", "Render or redirect"])
grid(["Situation", "System response"], [
    ["No active session", "Redirect to /login."],
    ["Signed in as a job seeker", "Redirect to /dashboard."],
    ["Signed in as a recruiter", "Redirect to /recruiter/jobs."],
    ["Signed in as an administrator", "Render the requested admin module."],
    ["Authentication state still loading", "Show a loading spinner and delay the route decision."],
], [5.2, 11.4])
info("Security principle:", "The React route guard improves navigation and user experience, but sensitive backend operations must also verify the caller. The manage-role Edge Function performs its own token and admin-role checks before using service-role privileges.", "DCFCE7")

heading("4. Shared admin interface and UX")
grid(["Interface area", "Behavior"], [
    ["Shared dashboard shell", "DashboardLayout gives the admin the same responsive sidebar, top bar, profile identity and sign-out behavior used by other workspaces."],
    ["Admin identity", "The sidebar shows the JobAI Scout brand and an Admin Portal label with a role-specific accent."],
    ["Navigation", "Admin Dashboard, Manage Users, Manage Jobs and Platform Analytics are selected according to the current route."],
    ["Responsive content", "Cards, tables and charts use responsive grids and horizontal overflow where large datasets cannot fit on a narrow screen."],
    ["Feedback", "Loading indicators prevent blank-state confusion; toast messages report success or failure after mutations."],
    ["Visual hierarchy", "Dark navy surfaces, violet/indigo accents, cards, badges and semantic status colors distinguish information from actions."],
], [4.1, 12.5])

heading("5. Platform overview dashboard")
para("AdminDashboard loads four datasets in parallel so the first screen can describe the health and scale of the product without waiting for one request at a time.")
grid(["Metric", "Source", "Calculation"], [
    ["Total Users", "profiles", "Exact row count using a head-only query."],
    ["Total Jobs", "jobs", "Exact row count across stored listings."],
    ["Applications", "job_applications", "Exact number of recruiter-side applications."],
    ["Total Fill Clicks", "extension_usage", "Number of loaded usage records, currently limited to the latest query result of at most 1,000 rows."],
    ["Total Fields Filled", "extension_usage.field_count", "Sum of every available field_count value."],
    ["Field breakdown", "extension_usage.fields", "Counts how often each semantic field name appears."],
    ["Top users", "extension_usage.email", "Groups clicks and filled-field totals by email, sorts by clicks, and shows the first ten."],
], [3.6, 4.8, 8.2])
info("Interpretation:", "The extension dashboard measures recorded fill activity, not successful job submissions. JobAI deliberately does not submit external applications for the user.")

heading("6. User management")
flow(["Load profiles", "Load roles", "Load applications", "Merge by user_id", "Display readiness", "Edit or change role", "Refresh list"])
grid(["Capability", "How it works"], [
    ["Unified user table", "Profiles, user_roles and job_applications are loaded concurrently and joined in the browser by user_id."],
    ["Profile visibility", "The table exposes identity/contact information, professional links, skills, desired roles, application count and profile completion indicators."],
    ["Edit profile", "The dialog edits full name, email, phone, biography, LinkedIn, GitHub, skills, desired roles and experience years."],
    ["Array normalization", "Comma-separated skills and desired roles are trimmed and saved as arrays; blank optional values become null."],
    ["Role management", "An administrator can promote a standard user to admin or demote an admin to user."],
    ["Self-protection", "The client blocks changing the current admin's own role, and the backend independently blocks self-demotion."],
], [4.2, 12.4])

heading("7. Secure role-change flow")
flow(["Admin clicks Promote/Demote", "Send access token", "manage-role validates JWT", "Verify caller in user_roles", "Validate target + role", "Service role updates row", "Reload users"])
grid(["Backend check", "Reason"], [
    ["Authorization header must contain a Bearer token", "Rejects anonymous callers with HTTP 401."],
    ["Token claims must resolve to a real user", "Prevents malformed or expired sessions from reaching privileged logic."],
    ["Caller must have admin role", "Rejects authenticated non-admin users with HTTP 403."],
    ["Target and new role are validated", "Only admin and user are accepted by this function."],
    ["Caller cannot demote themselves", "Reduces accidental lockout of the active administrator."],
    ["Service role is server-side only", "The browser never receives the privileged Supabase key."],
], [6.0, 10.6])

heading("8. Job and collection-source management")
grid(["Admin action", "Data behavior"], [
    ["Load listings", "Reads jobs in newest-first order."],
    ["Load sources", "Reads job_sources in newest-first order."],
    ["Add source", "Creates an enabled RSS/XML or official company-careers source after requiring a name and URL."],
    ["Enable or disable source", "Toggles job_sources.enabled without deleting its history."],
    ["Inspect source health", "Shows its URL, last error, last collection time and last result count when available."],
    ["Delete source", "Removes the selected source and reloads the source/listing view."],
    ["Delete job", "Removes one jobs row and immediately filters it out of local UI state."],
], [4.4, 12.2])
info("Operational caution:", "Source and job deletion are permanent database mutations in the current interface. The administrator should verify the target before using the trash action; the present implementation does not show a confirmation dialog.", "FEE2E2")

heading("9. Platform analytics")
para("AdminAnalytics currently renders the shared Analytics component. Because the administrator is neither a recruiter nor an ordinary seeker branch in its role-specific workflow, the screen primarily behaves as a platform job-market analytics view.")
grid(["Insight", "Calculation"], [
    ["Active jobs", "Reads up to 1,000 active jobs, optionally limited by the selected date range."],
    ["Location demand", "Groups listings by the first normalized location segment."],
    ["Top skills", "Counts skills attached to active job records and ranks the most frequent."],
    ["Job domains", "Classifies job title, description and skills using domain patterns."],
    ["Remote vs on-site", "Uses listing text and location indicators to estimate work arrangement."],
    ["Experience demand", "Places listings into experience buckets based on their text."],
    ["Posting trend", "Counts newly created jobs in date intervals for the selected range."],
    ["Top companies", "Groups and ranks active listings by company."],
], [4.4, 12.2])
info("Current scope:", "This route is not yet a dedicated admin-only business-intelligence implementation. It is a reuse of the general analytics page, which provides useful market insight but does not yet add admin-specific filters, exports or audit metrics.", "FEF3C7")

heading("10. Voice Assistant Control Room")
para("The Voice Control Room is the most technical admin module. It connects the product's voice conversations to knowledge retrieval evidence, operational telemetry, indexed sources and configurable defaults.")
grid(["Tab", "Purpose"], [
    ["Overview", "Shows conversations, messages, searches, failures, low-confidence retrievals, average confidence, latency, query activity and language mix."],
    ["Document Ingestion", "Lists indexed knowledge sources with source key, document type, status, page/chunk information and last crawl time."],
    ["Knowledge Search Inspector", "Runs a test retrieval and displays ranked knowledge chunks, similarity, source, page and section metadata."],
    ["System Config", "Controls assistant availability, silence timeout, RAG confidence threshold, default persona and default speaking speed."],
], [4.2, 12.4])

heading("11. Voice analytics data flow")
flow(["Read session", "Load global settings", "Load kb_sources", "Load latest 20 logs", "Count conversations/messages", "Aggregate logs", "Render charts + tables"])
grid(["Value", "How it is derived"], [
    ["Total conversations", "Exact count from voice_conversations."],
    ["Total messages", "Exact count from voice_messages."],
    ["Searches and failures", "All loaded voice_search_logs are counted and split by was_successful."],
    ["Low confidence", "Current UI marks stored confidence below 0.65 as low."],
    ["Average confidence", "Mean of non-null confidence_score values."],
    ["Average latency", "Mean of non-null response_latency_ms values."],
    ["Top queries", "Lowercases and groups equivalent query text, then selects the top five."],
    ["Daily and language charts", "Groups logs by created date and detected language."],
], [4.2, 12.4])

heading("12. Voice knowledge and configuration")
grid(["Control", "Admin meaning"], [
    ["Assistant enabled", "Intended to make voice availability active or inactive globally."],
    ["Silence timeout", "Sets the intended quiet period between 1 and 5 seconds before a spoken turn finishes."],
    ["Grounding threshold", "Controls how strong a RAG match should be before retrieved knowledge is trusted."],
    ["Default persona", "Chooses professional, friendly, recruiter or support-agent behavior."],
    ["Default speed", "Sets the default voice playback rate from 0.75x to 1.5x."],
    ["Indexed sources", "Provides evidence about which documents entered the retrieval system and whether ingestion succeeded."],
    ["Search inspector", "Helps an administrator inspect the chunks a question can retrieve before judging answer quality."],
], [4.3, 12.3])
info("Grounding guidance:", "A higher threshold generally reduces weak context matches but can produce more no-answer cases. A lower threshold increases recall but also raises the risk of irrelevant evidence. The interface recommends keeping it above 60%.", "FEF3C7")

heading("13. Error handling and system states")
grid(["State", "Visible response"], [
    ["Initial load", "A spinner remains visible while settings, sources, logs and statistics are prepared."],
    ["Read failure", "A destructive toast reports that the admin panel could not load."],
    ["Mutation in progress", "The affected button shows a loader or disables itself to reduce duplicate actions."],
    ["Successful change", "A success toast confirms the action, and data is reloaded where needed."],
    ["Empty dataset", "Tables and charts show a readable empty state instead of an unexplained blank panel."],
    ["Role-change failure", "The backend message is surfaced through a destructive toast."],
], [4.0, 12.6])

heading("14. Current implementation notes and improvement priorities")
para("The following items are important for a technically accurate understanding of the current code. They are improvement opportunities, not features that should be described as already complete.")
grid(["Current observation", "Recommended improvement"], [
    ["Voice Control Room is absent from adminNav.", "Add a visible Voice Operations sidebar item linked to /admin/voice."],
    ["AdminVoice posts to /voice-settings without action=global.", "Send ?action=global and verify the response before claiming global settings were updated."],
    ["The retrieval inspector sends a zero-filled placeholder embedding.", "Generate a real embedding through a protected backend function before calling hybrid_search_kb."],
    ["Clear Logs does not inspect the returned Supabase error.", "Check the response, report failures honestly and add a confirmation dialog."],
    ["Job/source deletion has no confirmation step.", "Use a confirmation dialog with the exact target name and destructive-action copy."],
    ["Dashboard extension usage is limited to 1,000 records.", "Move aggregation to a database RPC/view so totals remain accurate at scale."],
    ["Admin analytics reuses general market analytics.", "Add admin KPIs, date filters, cohort metrics, audit history and export capability."],
    ["Some reads rely on RLS rather than a single admin API boundary.", "Keep strict database policies and consider purpose-built admin RPCs for consistent auditing."],
], [7.3, 9.3])

heading("15. Recommended daily admin workflow")
grid(["Step", "Administrator task", "Reason"], [
    ["1", "Open Platform Overview and scan scale plus extension-usage indicators.", "Quickly detects unusual drops, spikes or missing telemetry."],
    ["2", "Review users and profile completeness; make role changes only when authorized.", "Maintains access integrity and user-data quality."],
    ["3", "Inspect job-source errors and collection recency before deleting listings.", "Separates ingestion failures from ordinary listing lifecycle changes."],
    ["4", "Use Platform Analytics to review market demand, locations and skills.", "Provides context for product and data-source decisions."],
    ["5", "Open /admin/voice to inspect failed or low-confidence searches and source health.", "Connects poor assistant outcomes to retrieval evidence."],
    ["6", "Test representative questions in the search inspector after proper embedding support is enabled.", "Verifies that the knowledge base retrieves relevant evidence."],
    ["7", "Change global configuration carefully and validate it with a real user session.", "Confirms that operational changes affect the intended runtime behavior."],
], [1.0, 8.0, 7.6])

heading("16. Reader summary")
para("JobAI Scout's admin panel is a role-protected operational layer over Supabase data and Edge Functions. Its overview combines platform counts with extension usage. User management joins profiles, roles and applications, while sensitive role changes pass through a server-verified function. Job management controls collection sources and listings. Analytics explains the active job market. The Voice Control Room exposes conversations, retrieval telemetry, indexed knowledge and intended global defaults. The design provides a strong foundation, while the documented navigation, settings, search-inspector, confirmation and aggregation improvements are the next steps required for a fully mature competition-grade administration system.")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("JobAI Scout Admin Panel - Complete Technical Explanation")
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor.from_string(SLATE)

doc.core_properties.title = "How the JobAI Scout Admin Panel Works"
doc.core_properties.author = "JobAI Scout"
doc.core_properties.subject = "Admin access, platform management, analytics and voice/RAG operations"
doc.save(OUT)
print(OUT)
