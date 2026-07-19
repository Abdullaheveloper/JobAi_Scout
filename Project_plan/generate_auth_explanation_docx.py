"""Create a reader-friendly technical explanation of JobAI Scout authentication."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(__file__).with_name("JobAI_Scout_Authentication_Explained.docx")
NAVY = "0B1028"
INDIGO = "4F46E5"
LIGHT = "EEF2FF"
SLATE = "334155"


def shade(cell, value):
    props = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), value)
    props.append(fill)


def border_bottom(paragraph):
    p = paragraph._p
    ppr = p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), INDIGO)
    borders.append(bottom)
    ppr.append(borders)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.7)
sec.bottom_margin = Cm(1.7)
sec.left_margin = Cm(1.8)
sec.right_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10)
normal.font.color.rgb = RGBColor.from_string("172033")
normal.paragraph_format.space_after = Pt(6)

for name, size, color in [("Title", 25, NAVY), ("Heading 1", 17, NAVY), ("Heading 2", 13, INDIGO), ("Heading 3", 11, VIOLET if (VIOLET := "7C3AED") else INDIGO)]:
    style = doc.styles[name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)

code = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
code.font.name = "Consolas"
code.font.size = Pt(8.5)
code.font.color.rgb = RGBColor.from_string(NAVY)
code.paragraph_format.space_after = Pt(8)


def title(text, subtitle=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(27)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    if subtitle:
        s = doc.add_paragraph(subtitle)
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.runs[0].font.color.rgb = RGBColor.from_string(SLATE)
        s.runs[0].font.size = Pt(11)


def h(text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        border_bottom(p)
    return p


def para(text, bold_start=None):
    p = doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        p.add_run(bold_start).bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def note(label, text, color="EEF2FF"):
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    shade(c, color)
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = c.paragraphs[0]
    p.add_run(label + " ").bold = True
    p.add_run(text)
    doc.add_paragraph()


def grid(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, value in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = value
        shade(c, NAVY)
        for r in c.paragraphs[0].runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if len(t.rows) % 2 == 0:
                shade(cells[i], "F8FAFC")
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return t


def flow(steps):
    t = doc.add_table(rows=1, cols=len(steps))
    t.autofit = False
    for i, text in enumerate(steps):
        c = t.cell(0, i)
        shade(c, LIGHT if i % 2 == 0 else "EDE9FE")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph()


title("How Authentication Works in JobAI Scout", "A practical, code-based explanation for readers, evaluators and future developers")
doc.add_paragraph()
note("Purpose:", "This document explains the authentication system as it is currently implemented in the JobAI Scout repository. It distinguishes user-interface protection from database security and identifies one role-assignment issue that should be fixed before production release.")

h("1. The simple idea", 1)
para("Authentication answers: “Who is this person?” Authorization answers: “What is this authenticated person allowed to do?” JobAI Scout uses Supabase Authentication to prove identity, then uses a separate role record and database policies to decide which screens and records each person may access.")
flow(["1. Register or sign in", "2. Supabase creates / verifies identity", "3. Session token is saved", "4. App loads profile + role", "5. Route and database enforce access"])
grid(["Layer", "What it does", "Main implementation"], [
    ["Identity", "Creates accounts, checks passwords, issues sessions and supports confirmation flows.", "Supabase Auth via supabase.auth.signUp and signInWithPassword."],
    ["Client session", "Keeps the signed-in state across refreshes and refreshes tokens.", "createClient auth storage: localStorage, persistSession: true, autoRefreshToken: true."],
    ["Application role", "Classifies the account as user, recruiter or admin.", "public.user_roles table; AuthContext loads the role."],
    ["User experience", "Shows the appropriate page or redirects the user.", "ProtectedRoute and React Router."],
    ["Real security", "Blocks data reads/writes that the current user is not allowed to perform.", "Supabase Row Level Security (RLS) policies and protected Edge Functions."],
], [3.2, 6.4, 7.0])

h("2. Where the authentication code lives", 1)
grid(["File / database object", "Responsibility"], [
    ["src/integrations/supabase/client.ts", "Creates the Supabase client and enables persistent local session storage plus automatic token refresh."],
    ["src/pages/Register.tsx", "Validates registration input, selects job-seeker or recruiter role, calls Supabase sign-up and handles email-confirmation state."],
    ["src/pages/Login.tsx", "Validates sign-in input, calls password sign-in and routes the user after a successful session response."],
    ["src/contexts/AuthContext.tsx", "Keeps user, session, loading state, role, profile and recruiter profile available across the React application."],
    ["src/components/ProtectedRoute.tsx", "Prevents unauthenticated or wrong-role users from seeing protected page components."],
    ["auth.users", "Supabase-managed identity table. It stores account identity and user metadata; frontend code does not write directly to it."],
    ["profiles, user_roles, recruiter_profiles", "Application data tables created and populated after a new authentication user exists."],
    ["handle_new_user() trigger", "Database function that creates the matching profile/role records immediately after signup."],
], [5.5, 11.1])

h("3. Registration: what happens when a new person creates an account", 1)
para("The registration page offers two ordinary account choices: job seeker (stored as user) and recruiter. It first performs lightweight client-side checks. These checks improve the user experience, but they are not the final security boundary because a malicious browser can bypass JavaScript. The backend must remain authoritative.")
grid(["Step", "What the code does", "Why it exists"], [
    ["1. Validate", "Checks full name, email pattern, password length of at least eight characters, and company name for recruiters.", "Catches common mistakes before an API call."],
    ["2. Build metadata", "Creates raw user metadata: full_name, role, and company_name when recruiter is selected.", "Provides initialization details for the database trigger."],
    ["3. Create identity", "Calls supabase.auth.signUp with email, password, metadata and emailRedirectTo set to /login.", "Supabase hashes and stores the password; the application never stores the password itself."],
    ["4. Confirm or enter", "If Supabase returns a session, the UI continues to the selected workspace. If no session is returned, it redirects to Login and tells the user to confirm email.", "Supports both immediate-session and email-confirmation configurations."],
    ["5. Database trigger", "After auth.users receives the account, handle_new_user() inserts profile, role and, for recruiters, recruiter profile rows.", "Creates the application records that the rest of the app expects."],
], [1.4, 8.8, 6.4])

note("Important distinction:", "The frontend sends the password only to Supabase Auth over HTTPS. JobAI Scout does not insert passwords into profiles, user_roles, local tables, logs or browser extension storage.", "DCFCE7")

h("4. The database trigger: automatic profile creation", 1)
para("A database trigger is a rule that runs automatically when a database event occurs. In this project, the on_auth_user_created trigger runs after an account is inserted into Supabase’s auth.users table. It calls public.handle_new_user(). The function runs with SECURITY DEFINER because it needs permission to create corresponding records in public tables even though the new user has not manually inserted them.")
flow(["auth.users INSERT", "on_auth_user_created trigger", "handle_new_user()", "profiles INSERT", "user_roles INSERT", "recruiter_profiles INSERT when needed"])
grid(["Record created", "Key data", "Purpose"], [
    ["profiles", "user_id, email, full_name", "Holds the user’s career profile data; user_id links it one-to-one to the authenticated user."],
    ["user_roles", "user_id, role", "Stores authorization role separately from general profile content."],
    ["recruiter_profiles", "user_id, company_name", "Stores recruiter-specific information only when the selected role is recruiter."],
], [4.1, 5.4, 7.1])

h("5. Sign-in and session restoration", 1)
para("On the Login page, the user enters email and password. The page normalizes the email to lowercase and calls supabase.auth.signInWithPassword. On a successful response, Supabase provides a session. A session includes access credentials that prove the user is signed in for a limited period. The client library stores the session in localStorage, restores it after page refresh, and automatically refreshes a token when possible.")
grid(["Moment", "AuthContext behavior"], [
    ["Application starts / refreshes", "Calls supabase.auth.getSession(). If a session exists, it sets user and session and loads profile and role."],
    ["Sign-in, sign-out, token refresh", "Subscribes through supabase.auth.onAuthStateChange. Every state change updates React state."],
    ["Signed-in user found", "fetchUserData(user.id) loads profiles and user_roles in parallel. If role is recruiter, it also loads recruiter_profiles."],
    ["No user", "Clears profile, role and recruiter profile state so the UI cannot display a prior account’s data."],
    ["User presses Sign out", "Calls supabase.auth.signOut(), clears local React state, and causes protected routes to send the person to Login."],
], [5.1, 11.5])
note("Why loading exists:", "ProtectedRoute waits while AuthContext is still checking the session. Without this waiting state, a valid user could be briefly redirected to Login during a page refresh before the saved session is restored.")

h("6. Roles and protected routes", 1)
para("JobAI Scout currently uses three application roles: user (job seeker), recruiter and admin. The frontend makes navigation friendly by using ProtectedRoute. For example, /dashboard/assistant requires user, recruiter pages require recruiter, and /admin pages require admin. If the current role is different, ProtectedRoute redirects to that account’s own fallback dashboard.")
grid(["Role", "Primary route group", "Typical capabilities"], [
    ["user", "/dashboard/*", "Career profile, CV upload, job browsing, saved jobs, applications, analytics, auto-fill and voice assistant."],
    ["recruiter", "/recruiter/*", "Recruiter profile, job publishing, candidate review and application-status workflow."],
    ["admin", "/admin/*", "User management, job management, platform analytics and administration tools."],
], [3.0, 4.2, 9.4])
note("Do not rely on routes alone:", "A route guard only changes what the browser renders. A determined user can still send a direct database or API request. The real protection is the Row Level Security policy evaluated by Supabase for every authenticated request.", "FEF3C7")

h("7. Row Level Security: the actual data protection", 1)
para("Row Level Security (RLS) is enabled on important application tables. Each policy uses auth.uid(), which is the user ID decoded from the verified Supabase access token. This means a browser cannot simply choose another user_id in a request and read that person’s private rows if the policy correctly compares the row owner to auth.uid().")
grid(["Table / resource", "Examples of the implemented rule"], [
    ["profiles", "A user can select, insert and update only a profile whose user_id equals auth.uid(); admins can view all profiles through has_role(auth.uid(), 'admin')."],
    ["user_roles", "A user can view own role; administrative role lookups are enabled by has_role()."],
    ["saved_jobs", "A user can view, save and remove only saved-job rows with their own user_id."],
    ["job_applications", "A job seeker reads own applications; administrators can view all; recruiter policies support job-owner workflows."],
    ["recruiter_profiles", "Recruiters manage their own profile; administrators have management access; authenticated users may view public recruiter profile information."],
    ["resume storage", "Storage policy limits file paths so the first folder component must equal the authenticated user’s ID."],
], [4.7, 11.9])
para("The has_role(user_id, role) database function is marked SECURITY DEFINER and checks whether a matching role record exists. It is used inside policies so the database can make a centralized authorization decision. This is appropriate only when the role table itself cannot be manipulated by ordinary users.")

h("8. How protected Edge Functions use authentication", 1)
para("Some operations must never be performed only in browser code because they depend on secret provider keys or protected business logic. JobAI Scout places these operations in Supabase Edge Functions, including CV analysis, job collection, document knowledge-base ingestion, cover-letter generation, voice transcription, voice chat and text-to-speech. The client obtains the active session and sends its access token in an Authorization: Bearer header. The function can verify the caller before accessing private data or making an AI/provider request.")
flow(["Browser has session", "Client gets access token", "Bearer token sent to Edge Function", "Function verifies caller", "Function reads permitted data / calls provider", "Sanitized response returns to browser"])
note("Secret rule:", "VITE_ variables are compiled into browser code and therefore are not secret. Only the Supabase URL and publishable key belong there. API keys for AI, job providers, service role access, or other privileged services must stay in Supabase secrets or another server-only secret store.", "FEE2E2")

h("9. Current security finding: role assignment must be hardened", 1)
para("The current trigger uses the value in NEW.raw_user_meta_data->>'role' and casts it to the app_role enum. The registration page normally offers only user and recruiter, but browser code is not a trust boundary. A person who calls the Supabase sign-up API directly may be able to submit a different allowed enum value. Because handle_new_user() runs as SECURITY DEFINER, this deserves immediate review before the system is presented as production secure.")
grid(["Current behavior", "Risk", "Safer design"], [
    ["Trigger accepts metadata role directly.", "If admin is an allowed app_role value, an attacker may attempt to request admin at signup.", "Always insert 'user' for public sign-up; never create admin from browser-provided metadata."],
    ["Recruiter is chosen at registration.", "Recruiter signup may be acceptable, but should be a deliberate product decision.", "Whitelist only user and recruiter in the trigger; reject every other value. Optionally require admin approval before recruiter privileges activate."],
    ["Admin role can be set by a migration/manual update.", "A hard-coded migration is not a scalable admin-provisioning process.", "Use a restricted, server-only admin management function or Supabase dashboard procedure with audit logging."],
], [4.8, 5.7, 6.1])
note("Recommended rule:", "Never trust raw_user_meta_data for elevated privileges. Metadata can describe a user, but a privileged role must be assigned by a trusted administrator or server-side process after verification.", "FEE2E2")

h("10. End-to-end examples", 1)
h("Example A: Job seeker creates an account", 2)
para("The user selects Job seeker, enters name, email and an eight-character password, and submits the Register form. Supabase creates auth.users. The trigger creates profiles and a user_roles record with user. If email confirmation is enabled, the UI directs the person to confirm their email before login. After sign-in, AuthContext loads user data, and ProtectedRoute permits /dashboard/* pages. When the user saves a job, the saved_jobs policy checks that auth.uid() equals that saved row’s user_id.")
h("Example B: Recruiter signs in and publishes a job", 2)
para("The recruiter account has recruiter role and recruiter_profiles data. AuthContext loads both. ProtectedRoute allows /recruiter/jobs. When the recruiter inserts a job, the database policy requires that auth.uid() equals recruiter_id and that has_role(auth.uid(), 'recruiter') is true. This creates two checks: the job must belong to the caller and the caller must hold recruiter role.")
h("Example C: Unauthenticated person opens a private route", 2)
para("When no session exists, AuthContext finishes loading with user equal to null. ProtectedRoute redirects the visitor to /login. Even if the person bypasses the interface and tries to query a private table directly, RLS policies should deny access because auth.uid() has no authenticated identity.")

h("11. Production checklist", 1)
for text in [
    "Enable email confirmation and configure the exact production redirect URL in Supabase Authentication settings.",
    "Replace direct metadata-based role assignment with a strict whitelist and server-controlled admin provisioning.",
    "Verify RLS is enabled and tested for every table, storage bucket and new feature before deployment.",
    "Keep the Supabase service-role key, AI keys and job-provider tokens out of .env values exposed to Vite and out of Git history.",
    "Use HTTPS in production; microphone and secure-session features depend on a trusted origin.",
    "Test signup, wrong password, email confirmation, token refresh, logout, wrong role, expired session and cross-account access using two test users.",
    "Add rate limiting / bot protection and audit logging for sensitive actions such as role changes, job collection and admin management.",
]:
    bullet(text)

h("12. Reader summary", 1)
para("In one sentence: Supabase proves who the user is, AuthContext keeps that session and role available to React, ProtectedRoute gives the correct interface, and RLS plus authenticated Edge Functions must enforce what the user is actually allowed to read or change. The most important improvement before public production use is to make privileged role assignment server-controlled rather than trusting signup metadata.")

section = doc.sections[0]
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("JobAI Scout Authentication - Technical Explanation")
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor.from_string(SLATE)

doc.core_properties.title = "How Authentication Works in JobAI Scout"
doc.core_properties.author = "JobAI Scout"
doc.save(OUT)
print(OUT)
