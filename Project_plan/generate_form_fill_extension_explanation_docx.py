"""Generate the complete Form Fill tab and browser extension explanation."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name("JobAI_Scout_Form_Fill_and_Extension_Explained.docx")
NAVY, INDIGO, VIOLET, SLATE = "0B1028", "4F46E5", "7C3AED", "334155"


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), color)
    props.append(node)


def bottom_border(paragraph):
    props = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), INDIGO)
    border.append(bottom)
    props.append(border)


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Cm(1.65)
section.left_margin = section.right_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Aptos", Pt(10)
normal.font.color.rgb = RGBColor.from_string("172033")
normal.paragraph_format.space_after = Pt(6)
for name, size, color in [("Heading 1", 17, NAVY), ("Heading 2", 13, INDIGO), ("Heading 3", 11, VIOLET)]:
    style = doc.styles[name]
    style.font.name, style.font.size, style.font.bold = "Aptos Display", Pt(size), True
    style.font.color.rgb = RGBColor.from_string(color)


def heading(text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        bottom_border(paragraph)
    return paragraph


def para(text):
    return doc.add_paragraph(text)


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def info(label, text, color="EEF2FF"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, color)
    paragraph = cell.paragraphs[0]
    paragraph.add_run(label + " ").bold = True
    paragraph.add_run(text)
    doc.add_paragraph()


def grid(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        shade(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.bold, run.font.color.rgb = True, RGBColor(255, 255, 255)
    for row_number, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = str(text)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_number % 2 == 0:
                shade(cells[index], "F8FAFC")
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()


def flow(steps):
    table = doc.add_table(rows=1, cols=len(steps))
    for index, text in enumerate(steps):
        cell = table.cell(0, index)
        shade(cell, "EEF2FF" if index % 2 == 0 else "EDE9FE")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold, run.font.size = True, Pt(8)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph()


cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
title = cover.add_run("How Form Fill and the JobAI Browser Extension Work")
title.font.name, title.font.size, title.font.bold = "Aptos Display", Pt(25), True
title.font.color.rgb = RGBColor.from_string(NAVY)
subtitle = doc.add_paragraph("Complete user, technical, security and decision-flow explanation")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.color.rgb, subtitle.runs[0].font.size = RGBColor.from_string(SLATE), Pt(11)
doc.add_paragraph()
info("Purpose:", "This document explains the Form Fill dashboard tab and its Chrome-compatible extension as one connected product. It covers profile preparation, installation, authentication, field detection, confidence rules, resume attachment, privacy, limitations, cleanup and troubleshooting.")

heading("1. The feature in one sentence")
para("Form Fill converts career facts that the user has deliberately saved in the JobAI Scout Career Passport into careful assistance on external job-application forms. The browser extension fills reliable facts, suggests uncertain matches for review, and leaves sensitive, legal and final decisions to the applicant.")
flow(["Build Career Passport", "Download extension", "Install and sign in", "Open application", "Detect fields", "Fill or flag", "User reviews and submits"])
grid(["Part", "Responsibility", "Primary implementation"], [
    ["Form Fill tab", "Explains behavior, safety levels, supported profile data and installation; provides the canonical extension download.", "src/pages/AutoFormFill.tsx"],
    ["Extension popup", "Authenticates the user, displays Career Passport readiness, uploads a resume and starts a fill pass.", "extension/popup.html and popup.js"],
    ["Profile service", "Loads and caches the signed-in user's normalized profile through the protected extension API.", "extension/profile-service.js and api.js"],
    ["Decision engine", "Classifies fields, applies confidence and evidence rules, and blocks protected decisions.", "extension/decision-engine.js"],
    ["Page engine", "Finds controls, resolves values, fills supported controls and displays the review panel.", "extension/content.js"],
], [3.2, 7.5, 5.9])

heading("2. What the Form Fill tab provides")
grid(["Interface area", "Reader-friendly explanation"], [
    ["Premium introduction", "Explains that the product saves application effort without taking judgment away from the user."],
    ["Download extension", "Downloads the single supported package, job-form-fill.zip, with visible loading and failure feedback."],
    ["Build Career Passport", "Links directly to Profile Settings so missing facts can be completed before filling forms."],
    ["Decision model", "Separates automatic factual filling, suggestions for review and manual-only decisions."],
    ["Field coverage", "Shows identity, contact, education, work history, links, skills, preferences and document behavior."],
    ["Installation workflow", "Presents four calm steps from profile preparation through final user review."],
], [4.2, 12.4])
info("Naming decision:", "The navigation label is Form Fill. The route remains /dashboard/auto-fill for backward compatibility, so saved links do not break.", "DCFCE7")

heading("3. Career Passport: the source of truth")
para("The extension should not invent personal facts. Its strongest answers come from Profile Settings and CV review, where the user can approve or correct information before reuse. The profile can contain contact details, education, employment history, projects, certifications, languages, professional links, application preferences and saved screening answers.")
grid(["Profile group", "Examples used during form filling"], [
    ["Identity and contact", "Full name, first and last name, email, phone, city, location and country."],
    ["Professional links", "LinkedIn, GitHub and portfolio URLs."],
    ["Education", "Institution, degree, subject, start date, graduation/end date and achievements."],
    ["Work history", "Employer, title, dates, duties, achievements and calculated experience years."],
    ["Career evidence", "Skills, projects, certifications, languages and verified profile facts."],
    ["Preferences", "Work authorization, relocation and commute answers only when explicitly saved and safe to reuse."],
    ["Documents", "A private resume path used for supported resume upload controls."],
], [4.2, 12.4])

heading("4. Installation and sign-in")
grid(["Step", "What the user does", "What the system does"], [
    ["1", "Open Form Fill and download job-form-fill.zip.", "The website retrieves the canonical package from the public deployment."],
    ["2", "Extract the ZIP into a permanent folder.", "Chrome requires an unpacked folder containing manifest.json and the extension files."],
    ["3", "Open chrome://extensions and enable Developer mode.", "The browser exposes Load unpacked and Reload controls."],
    ["4", "Choose Load unpacked and select the extracted folder.", "Chrome validates Manifest V3 permissions and starts the background worker."],
    ["5", "Open the extension and sign in with JobAI email/password or Google.", "Supabase returns a short-lived user session; refresh logic renews it when required."],
    ["6", "Refresh an already-open application tab after installation.", "The content script becomes available on an authorized job-application host."],
], [1.0, 7.2, 8.4])
info("Cross-laptop rule:", "Each laptop needs its own extracted extension folder and one-time Load unpacked installation. The JobAI account and Career Passport remain shared through the hosted backend, but browser installation is local to each device.", "FEF3C7")

heading("5. Complete field-processing flow")
flow(["Scan controls", "Read labels and context", "Classify semantic meaning", "Resolve profile evidence", "Apply safety policy", "Fill or review", "Show outcome panel"])
grid(["Stage", "How it works"], [
    ["Control discovery", "The content script scans visible text inputs, text areas, selects, content-editable controls, radio buttons, checkboxes and file inputs."],
    ["Context construction", "It reads type, autocomplete, name, id, placeholder, ARIA label, associated labels and nearby question text."],
    ["Semantic classification", "Patterns map different employer wording to concepts such as email, graduation date, employer, experience, commute or resume."],
    ["Value resolution", "The engine searches normalized direct profile fields, structured Career Passport collections and explicitly saved answers."],
    ["Decision", "The safety engine evaluates control type, protected categories, confidence and whether direct evidence exists."],
    ["Interaction", "Supported values are inserted and normal focus, input, change and blur events are dispatched so React-based forms detect them."],
    ["Review", "A side panel lists filled fields, suggestions, missing facts and protected controls. JobAI never presses the final submit button."],
], [4.0, 12.6])

heading("6. Confidence and safety rules")
grid(["Confidence / category", "Result", "Reason"], [
    ["Below 40%", "Leave blank", "The match is too weak even to present as a useful answer."],
    ["40% to 74%", "Suggestion only", "The user may review the possible mapping, but JobAI does not place it as a decision."],
    ["75% to 89%", "Low-risk text may fill with review", "Only ordinary text facts can be placed and must remain visible for confirmation."],
    ["90% or higher plus direct evidence", "Safe factual controls may fill", "High confidence alone is insufficient; the value must be backed by saved evidence."],
    ["Checkboxes and radio buttons", "Require at least 90% and direct, non-sensitive evidence", "A 40% checkbox match is a prompt, not permission to select an answer."],
    ["Legal, diversity, consent, CAPTCHA and submit", "Always manual", "These represent personal declarations, protected data, verification or final intent."],
], [4.5, 5.2, 6.9])
info("Applicant control:", "Terms, privacy consent, disability, veteran status, ethnicity, gender/self-identification, CAPTCHA, assessments, signatures and final submission remain manual regardless of any confidence score.", "FEE2E2")

heading("7. Resume upload and automatic attachment")
para("The extension popup includes an Application resume card. The user can upload or replace a PDF or DOCX of up to 10 MB. The file is stored in the private resumes bucket under that user's own ID, and only the private storage path is saved in the profile.")
flow(["Choose PDF/DOCX", "Validate type and size", "Private storage upload", "Update resume_url", "Refresh profile", "Detect resume input", "Download privately and attach"])
grid(["Behavior", "Implementation detail"], [
    ["Private upload", "The request carries the signed-in access token and configured public Supabase key; storage policies identify the user."],
    ["Failure safety", "Invalid or oversized files are rejected before upload. A failed upload does not replace the user's previous resume."],
    ["Correct file type", "The attachment logic preserves the real PDF or DOCX filename and MIME type so an ATS does not reject a DOCX renamed as PDF."],
    ["Standard ATS input", "The extension downloads the private file, builds a File and DataTransfer, assigns the file list and dispatches the expected events."],
    ["No stored resume", "The field is listed as missing rather than guessed or silently ignored."],
], [4.4, 12.2])

heading("8. Google Forms file-upload limitation")
para("A Google Forms 'Add file' question is not a normal file input available to the page. It opens Google's authenticated Drive picker, which runs behind a separate protected interface. The extension therefore detects this control and shows a clear manual-review message instead of claiming success.")
grid(["Google Forms behavior", "JobAI response"], [
    ["Ordinary text, date and supported choice fields", "May be filled when the mapping, confidence and evidence policy allow it."],
    ["Add file / Drive picker", "The user must click Add file and choose the resume through Google's signed-in picker."],
    ["Final submission", "The user reviews all answers and submits the form personally."],
], [5.0, 11.6])
info("Why this is correct:", "Attempting to bypass Google's protected picker would be unreliable and unsafe. Honest manual guidance is better UX than showing a false 'resume attached' result.", "FEF3C7")

heading("9. Supported browsers and application sites")
para("The extension uses Manifest V3 and is intended for Chrome, Microsoft Edge, Brave, Arc, Opera and other Chromium-based browsers that support unpacked extensions. Its manifest authorizes selected job application hosts, including Greenhouse, Lever, Workday, Ashby, SmartRecruiters, Jobvite, iCIMS, Workable, LinkedIn, Indeed and Google Forms. Unknown or unauthorized websites are not granted silent access.")
info("Scope:", "Form Fill assists with applications. Job discovery and scraping belong to Browse Jobs and server-side collection services; the extension does not pretend to be a job-scanning tool.", "DCFCE7")

heading("10. Authentication, privacy and security")
grid(["Control", "Protection"], [
    ["User authentication", "Email/password and Google sign-in use Supabase Auth. Session refresh prevents an expired token from being reused indefinitely."],
    ["Profile boundary", "The extension-profile function verifies the caller and returns only the profile belonging to that authenticated user."],
    ["Private resume", "Storage paths must begin with the signed-in user's ID; downloads and uploads include the user's access token."],
    ["Host permissions", "The build script synchronizes the exact configured Supabase host into manifest.json, preventing cross-laptop CORS/permission inconsistency."],
    ["Local cache", "Only the active session and normalized profile cache are stored in chrome.storage.local; logout removes both."],
    ["No submission", "The extension does not accept declarations or submit the employer's form."],
], [4.3, 12.3])

heading("11. User experience states")
grid(["State", "Visible behavior"], [
    ["Signed out", "A focused login panel offers email/password and Google sign-in with friendly errors."],
    ["Profile loaded", "The popup shows identity details, skills and Career Passport completion."],
    ["Resume missing / ready", "A badge and upload/replace action communicate document readiness."],
    ["Preparing", "Buttons are disabled and a loading indicator prevents multiple simultaneous fill requests."],
    ["Complete", "Filled, review and protected counts summarize the result."],
    ["Error", "The popup provides a readable action-oriented message instead of a raw technical failure."],
], [4.0, 12.6])

heading("12. Troubleshooting")
grid(["Problem", "Resolution"], [
    ["Extension changes do not appear", "Open chrome://extensions, press Reload on Job Form Fill, then refresh the application page."],
    ["Download opens but will not install", "Extract the ZIP first; Load unpacked must point to the folder containing manifest.json."],
    ["Session expired", "Open the popup and sign in again. Check internet access if refresh fails."],
    ["A known field stays empty", "Add or verify the fact in Profile Settings, refresh the Career Passport in the popup, and run Fill again."],
    ["Resume is not attached", "Confirm the popup says Ready, the form uses a standard file input, and the site's accepted type includes the saved file."],
    ["Google Forms asks for a file", "Click Add file and choose it manually through the Google Drive picker."],
    ["New laptop", "Download/extract the latest package and load it once on that browser; do not reuse an old extracted version."],
], [5.0, 11.6])

heading("13. Testing and cleanup design")
para("The extension has decision-engine safety tests, a content-script smoke test, Google Forms mapping regression tests and a resume-upload contract test. The production website build also verifies that Form Fill compiles with the rest of JobAI Scout.")
grid(["Test", "Coverage"], [
    ["Decision engine", "Confidence boundaries, direct evidence and manual-only categories."],
    ["Content smoke test", "Text, radio, checkbox and resume filling plus duplicate-pass protection."],
    ["Google Forms regression", "Graduation date, institution, experience, current employer and commute mapping."],
    ["Resume contract", "Popup controls, private upload path, profile update, DataTransfer attachment and protected Google picker behavior."],
    ["Production build", "TypeScript/React bundling and the canonical extension configuration sync."],
], [4.2, 12.4])
info("Single-extension cleanup:", "The obsolete Extension page and duplicate jobai-extension.zip were removed. The old /dashboard/extension address redirects to /dashboard/auto-fill, while job-form-fill.zip is now the only supported downloadable package.", "DCFCE7")

heading("14. Reader summary")
para("JobAI Scout Form Fill is an evidence-led application assistant rather than an uncontrolled bot. The dashboard tab teaches the workflow and provides one current download. The extension signs the user in, reads the approved Career Passport, uploads and privately reuses a resume, understands many employer field labels, applies transparent confidence rules, and produces a review panel. High-confidence factual work can be accelerated, while uncertain, sensitive, legal and final decisions remain visibly under the applicant's control.")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("JobAI Scout Form Fill & Extension - Complete Technical Explanation")
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor.from_string(SLATE)

doc.core_properties.title = "How Form Fill and the JobAI Browser Extension Work"
doc.core_properties.author = "JobAI Scout"
doc.core_properties.subject = "Form Fill dashboard tab, browser extension, safety rules and resume workflow"
doc.save(OUT)
print(OUT)
