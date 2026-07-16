"""CV/resume text extraction: PyMuPDF → pdfplumber → OCR fallback."""

from __future__ import annotations

import io
from typing import Any

import fitz  # PyMuPDF
import pdfplumber
from docx import Document

MIN_CHARS_PER_PAGE = 50
MAX_TEXT_LENGTH = 50000


def _clean_text(text: str) -> str:
    lines = [line.rstrip() for line in text.splitlines()]
    cleaned: list[str] = []
    blank_run = 0
    for line in lines:
        if not line.strip():
            blank_run += 1
            if blank_run <= 2:
                cleaned.append("")
            continue
        blank_run = 0
        cleaned.append(line)
    return "\n".join(cleaned).strip()


def _truncate(text: str, limit: int = MAX_TEXT_LENGTH) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[... truncated ...]"


def _extract_pdf_pymupdf(data: bytes) -> tuple[str, int]:
    doc = fitz.open(stream=data, filetype="pdf")
    pages: list[str] = []
    for index, page in enumerate(doc, start=1):
        text = page.get_text("text") or ""
        pages.append(f"[PAGE {index}]\n{text.strip()}")
    return "\n\n".join(pages), len(doc)


def _extract_pdf_pdfplumber(data: bytes) -> tuple[str, int]:
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages: list[str] = []
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            pages.append(f"[PAGE {index}]\n{text.strip()}")
        return "\n\n".join(pages), len(pdf.pages)


def _extract_pdf_ocr(data: bytes) -> tuple[str, int]:
    import pytesseract
    from PIL import Image

    doc = fitz.open(stream=data, filetype="pdf")
    pages: list[str] = []
    for index, page in enumerate(doc, start=1):
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img) or ""
        pages.append(f"[PAGE {index}]\n{text.strip()}")
    return "\n\n".join(pages), len(doc)


def _ocr_available() -> bool:
    try:
        import pytesseract

        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def extract_pdf(data: bytes) -> dict[str, Any]:
    text, page_count = _extract_pdf_pymupdf(data)
    method = "pymupdf"
    ocr_used = False
    avg_chars = len(text.strip()) / max(page_count, 1)

    if avg_chars < MIN_CHARS_PER_PAGE and page_count > 0:
        plumber_text, _ = _extract_pdf_pdfplumber(data)
        plumber_avg = len(plumber_text.strip()) / max(page_count, 1)
        if plumber_avg > avg_chars:
            text = plumber_text
            method = "pdfplumber"
            avg_chars = plumber_avg

        if avg_chars < MIN_CHARS_PER_PAGE and _ocr_available():
            try:
                text, _ = _extract_pdf_ocr(data)
                method = "ocr"
                ocr_used = True
            except Exception:
                pass

    text = _truncate(_clean_text(text))
    return {
        "text": text,
        "method": method,
        "pages": page_count,
        "ocr_used": ocr_used,
        "char_count": len(text),
    }


def extract_docx(data: bytes) -> dict[str, Any]:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    text = _truncate(_clean_text("\n".join(parts)))
    return {
        "text": text,
        "method": "python-docx",
        "pages": 1,
        "ocr_used": False,
        "char_count": len(text),
    }


def extract_file(data: bytes, filename: str) -> dict[str, Any]:
    lower_name = (filename or "").lower()

    if lower_name.endswith(".pdf"):
        return extract_pdf(data)
    if lower_name.endswith(".docx") or lower_name.endswith(".doc"):
        if lower_name.endswith(".doc"):
            # Legacy .doc is not supported by python-docx; try raw decode as last resort.
            try:
                return extract_docx(data)
            except Exception:
                text = _truncate(_clean_text(data.decode("utf-8", errors="ignore")))
                return {
                    "text": text,
                    "method": "raw-decode",
                    "pages": 1,
                    "ocr_used": False,
                    "char_count": len(text),
                }
        return extract_docx(data)

    text = _truncate(_clean_text(data.decode("utf-8", errors="ignore")))
    return {
        "text": text,
        "method": "plain-text",
        "pages": 1,
        "ocr_used": False,
        "char_count": len(text),
    }
